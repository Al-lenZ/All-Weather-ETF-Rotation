"""v6/leverage/report_docx_round_d.py — Chinese Word (.docx) report on the
higher-cap experiment (Round D vs Round C, both on the rep-set book).

Focus: the effect of raising the leverage cap 2.0 → 5.0 and the vol target
σ* 3.2 % → 6.4 %. Rep-set is mentioned briefly at the top. Internal
pre-registered pass gates are intentionally NOT surfaced. Duration risk
IS discussed.

Deliverables:
    reports/leverage_higher_cap_round_d_zh.docx  (Chinese Word report)

Run
---
    cd v6/leverage && python report_docx_round_d.py

Requires figures already produced by ``report_figures_round_d.py``.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import _common_leverage as CL


# =========================================================================
# constants
# =========================================================================
FIG_DIR = Path("/Users/allenzhou/Downloads/YSJ Lab/etf_basket_strategy/v6/reports/figures_round_d")
OUT_PATH = Path("/Users/allenzhou/Downloads/YSJ Lab/etf_basket_strategy/v6/reports/leverage_higher_cap_round_d_zh.docx")

# (label, control cap=2 cell, primary cap=5 cell)
PAIRS = [
    ("Base · GC007",       "C_base_reps_lev",       "D_base_reps_lev"),
    ("Base · DR007代理",   "C_base_reps_lev_DR007", "D_base_reps_lev_DR007"),
    ("EW · GC007（对照）",  "C_ew_reps_lev",         "D_ew_reps_lev"),
    ("EW · DR007代理（对照）", "C_ew_reps_lev_DR007", "D_ew_reps_lev_DR007"),
]

CELL_LABEL = {
    "C_base_reps_lev":       "Base · cap=2 · σ*=3.2% · GC007",
    "C_base_reps_lev_DR007": "Base · cap=2 · σ*=3.2% · DR007代理",
    "C_ew_reps_lev":         "EW · cap=2 · σ*=3.2% · GC007",
    "C_ew_reps_lev_DR007":   "EW · cap=2 · σ*=3.2% · DR007代理",
    "D_base_reps_lev":       "Base · cap=5 · σ*=6.4% · GC007",
    "D_base_reps_lev_DR007": "Base · cap=5 · σ*=6.4% · DR007代理",
    "D_ew_reps_lev":         "EW · cap=5 · σ*=6.4% · GC007",
    "D_ew_reps_lev_DR007":   "EW · cap=5 · σ*=6.4% · DR007代理",
}


# =========================================================================
# Chinese font helpers (copied from report_docx.py)
# =========================================================================
BODY_FONT_EAST = "宋体"
HEAD_FONT_EAST = "黑体"
BODY_FONT_ASCII = "Times New Roman"


def _set_run_font(run, size_pt: float, bold: bool = False,
                  color: RGBColor | None = None,
                  east: str = BODY_FONT_EAST) -> None:
    run.font.name = BODY_FONT_ASCII
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east)
    rFonts.set(qn("w:ascii"), BODY_FONT_ASCII)
    rFonts.set(qn("w:hAnsi"), BODY_FONT_ASCII)


def add_heading_zh(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 0:
        _set_run_font(r, 18, bold=True, east=HEAD_FONT_EAST)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        _set_run_font(r, 14, bold=True, east=HEAD_FONT_EAST)
    elif level == 2:
        _set_run_font(r, 12, bold=True, east=HEAD_FONT_EAST)
    else:
        _set_run_font(r, 11, bold=True, east=HEAD_FONT_EAST)


def add_para(doc, text: str, size: float = 10.5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size, bold=bold)


def add_bullet(doc, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_run_font(r, size)


# =========================================================================
# Table helpers
# =========================================================================
def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)


def add_table(doc, header: list[str], rows: list[list[str]],
              header_size: float = 9.5, body_size: float = 9.5,
              col_widths_cm: list[float] | None = None,
              highlight_rows: set[int] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        _set_run_font(r, header_size, bold=True, east=HEAD_FONT_EAST)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(c, "D9E1F2")

    highlight_rows = highlight_rows or set()
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[1 + i].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            _set_run_font(r, body_size, bold=(i in highlight_rows))
            c.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            )
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i in highlight_rows:
                _shade_cell(c, "FFF2CC")

    if col_widths_cm is not None:
        for j, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[j].width = Cm(w)


def add_figure(doc, fname: str, caption: str, width_cm: float = 16.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG_DIR / fname), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    _set_run_font(cr, 9, bold=False, color=RGBColor(0x40, 0x40, 0x40))


# =========================================================================
# data loaders
# =========================================================================
def load_summary(cell: str) -> pd.DataFrame:
    return pd.read_csv(CL.LEV_DIR / cell / "summary.csv").set_index("window")


def load_per_year(cell: str) -> pd.DataFrame:
    return pd.read_csv(CL.LEV_DIR / cell / "per_year.csv")


def load_net_ret(cell: str) -> pd.Series:
    return pd.read_csv(CL.LEV_DIR / cell / "net_ret.csv",
                       index_col=0, parse_dates=True)["net_ret"]


def load_lt(cell: str) -> pd.DataFrame:
    return pd.read_csv(CL.LEV_DIR / cell / "L_t_path.csv",
                       index_col=0, parse_dates=True)


def _pct(x, d: int = 2) -> str:
    return f"{x*100:+.{d}f}%" if pd.notna(x) else "—"


# ------------------------------------------------------------------ #
# Per-year DD on the CONTINUOUS NAV path + DD duration / recovery.
# ------------------------------------------------------------------ #
def per_year_dd_continuous(cell: str) -> pd.DataFrame:
    """Rebuild per-year metrics from the continuous NAV path so per-year
    MaxDD is directly comparable to window MaxDD (both use the same
    cummax-across-whole-window semantics). Also adds:

      - dd_peak_date   : last date NAV hit the running max before the
                         year's trough (i.e., where the DD started).
      - dd_trough_date : within-year date of MaxDD.
      - dd_duration_w  : weeks from peak to trough.
      - dd_recover_w   : weeks from trough until NAV recovers to the
                         prior peak; NaN if not recovered by end of the
                         cell's series (i.e., through 2025-07-31).

    Notes:
      - Window used is [first non-zero net_ret bar, 2025-07-31] — same
        as summary.csv's IS+OOS pooled window.
      - Additive-cumsum NAV is used everywhere (matches the summary
        convention). This means the multi-year peak that a within-year
        trough is measured against IS an early-window peak, not the
        year-open NAV.
    """
    nr = load_net_ret(cell)
    first_nz = nr[nr.ne(0.0)].index.min()
    mask = (nr.index >= first_nz) & (nr.index <= CL.OOS_END)
    nr = nr[mask]

    nav    = 1.0 + nr.cumsum()
    cummax = nav.cummax()
    dd     = (nav - cummax) / cummax
    L_ser  = load_lt(cell)["L_t"].reindex(nr.index)

    rows: list[dict] = []
    for yr, g in nr.groupby(nr.index.year):
        n = int(len(g))
        # Compound annualization, matches per_year.csv `cagr_net` (window_stats).
        cumret_year = float(g.sum())
        n_years     = max(n / 52.0, 1e-3)
        cagr_year   = max(1.0 + cumret_year, 1e-9) ** (1.0 / n_years) - 1.0
        L_year      = float(L_ser.loc[g.index].mean())
        # MaxDD on continuous NAV, restricted to this year's bars.
        dd_year = dd.loc[g.index]
        if len(dd_year) == 0 or dd_year.min() >= 0:
            rows.append(dict(year=int(yr), n=n, cagr=cagr_year, mean_L=L_year,
                             max_dd=0.0, dd_peak_date=None,
                             dd_trough_date=None,
                             dd_duration_w=None, dd_recover_w=None))
            continue
        trough_dt = dd_year.idxmin()
        max_dd    = float(dd_year.loc[trough_dt])
        peak_val  = float(cummax.loc[trough_dt])
        # find last date at or before trough where NAV == peak_val
        pre = nav.loc[:trough_dt]
        peak_dt = pre[pre >= peak_val - 1e-12].index.max()
        # weeks are the number of W-FRI bars between two dates on this series
        dur = int((nav.index.get_loc(trough_dt) - nav.index.get_loc(peak_dt)))
        # recovery: first bar after trough where NAV >= peak_val
        post = nav.loc[trough_dt:]
        rec_hits = post[post >= peak_val - 1e-12]
        if len(rec_hits) > 0:
            rec_dt = rec_hits.index[0]
            rec_w  = int((nav.index.get_loc(rec_dt) - nav.index.get_loc(trough_dt)))
        else:
            rec_w = None
        rows.append(dict(year=int(yr), n=n, cagr=cagr_year, mean_L=L_year,
                         max_dd=max_dd,
                         dd_peak_date=peak_dt.date(),
                         dd_trough_date=trough_dt.date(),
                         dd_duration_w=dur,
                         dd_recover_w=rec_w))
    return pd.DataFrame(rows).set_index("year")


# =========================================================================
# Sections
# =========================================================================
def build_summary_row(cell: str, window: str) -> list[str]:
    r = load_summary(cell).loc[window]
    return [
        CELL_LABEL[cell],
        f"{r['mean_L']:.2f}",
        _pct(r['cagr_net']),
        _pct(r['max_dd']),
        _pct(r['vol_realized']),
        f"{r['funding_drag_bp_yr']:+.0f}",
        f"{r['book_duration_yr_mean']:.1f}",
        f"{r['book_duration_yr_p95']:.1f}",
    ]


def build_summary_table(doc, window: str) -> None:
    header = ["组合", "平均 L̄", "净 CAGR", "最大回撤", "年化波动",
              "融资成本 (bp/年)", "久期均值 (年)", "久期 P95 (年)"]
    rows = []
    highlight = set()
    row_idx = 0
    for pair_label, c_cell, d_cell in PAIRS:
        rows.append(build_summary_row(c_cell, window))
        row_idx += 1
        rows.append(build_summary_row(d_cell, window))
        # highlight cap=5 (D) rows for base RB — the main line
        if d_cell in ("D_base_reps_lev", "D_base_reps_lev_DR007"):
            highlight.add(row_idx)
        row_idx += 1
    add_table(doc, header, rows,
              col_widths_cm=[4.8, 1.4, 1.7, 1.7, 1.6, 2.2, 1.9, 1.9],
              highlight_rows=highlight)


def _fmt_dur_cell(dur_w, rec_w) -> str:
    """Format DD duration cell: peak-to-trough weeks + trough-to-recovery."""
    if dur_w is None:
        return "—"
    rec_s = f"{rec_w}w" if rec_w is not None else "未回补"
    return f"{int(dur_w)}w → {rec_s}"


def _build_per_year_pair_table(doc, cells: list[tuple[str, str]]) -> None:
    """Two lines per year × per cell: CAGR + L̄, then DD (continuous NAV)
    with drawdown & recovery duration (weeks). MaxDD on the continuous
    NAV path so it is directly comparable with the summary table."""
    per_year_data = {cell: per_year_dd_continuous(cell) for _, cell in cells}
    all_years = sorted(set().union(*(set(pd.index) for pd in per_year_data.values())))
    header = ["年份 · 指标"] + [lbl for lbl, _ in cells]
    rows = []
    for y in all_years:
        # row 1: CAGR + L̄
        r1 = [f"{y} · CAGR / L̄"]
        r2 = [f"{y} · MaxDD / 时长（下跌→回补）"]
        for _, cell in cells:
            df = per_year_data[cell]
            if y not in df.index:
                r1.append("—"); r2.append("—"); continue
            row = df.loc[y]
            r1.append(f"CAGR {_pct(row['cagr'])} · L̄ {float(row['mean_L']):.2f}")
            dd_s = _pct(row['max_dd'])
            dur  = _fmt_dur_cell(row['dd_duration_w'], row['dd_recover_w'])
            r2.append(f"{dd_s} · {dur}")
        rows.append(r1); rows.append(r2)
    add_table(doc, header, rows,
              col_widths_cm=[2.8, 4.0, 4.0, 4.0, 4.0])


def build_per_year_table(doc) -> None:
    """Base RB per-year: CAGR / L̄ + DD (continuous) / duration."""
    cells = [
        ("Base cap=2 GC007（对照）", "C_base_reps_lev"),
        ("Base cap=5 GC007",         "D_base_reps_lev"),
        ("Base cap=2 DR007（对照）", "C_base_reps_lev_DR007"),
        ("Base cap=5 DR007",         "D_base_reps_lev_DR007"),
    ]
    _build_per_year_pair_table(doc, cells)


def build_per_year_table_ew(doc) -> None:
    """EW RB per-year (对照)."""
    cells = [
        ("EW cap=2 GC007（对照）",   "C_ew_reps_lev"),
        ("EW cap=5 GC007",           "D_ew_reps_lev"),
        ("EW cap=2 DR007（对照）",   "C_ew_reps_lev_DR007"),
        ("EW cap=5 DR007",           "D_ew_reps_lev_DR007"),
    ]
    _build_per_year_pair_table(doc, cells)


# =========================================================================
# main
# =========================================================================
def build_doc() -> None:
    doc = Document()

    # -- default styles
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT_ASCII
    style.font.size = Pt(10.5)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), BODY_FONT_EAST)

    # ============= Title =============
    add_heading_zh(doc, "v6 组合杠杆上限提升实验报告", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("对比 cap 5 σ*=6.4% 与 cap 2 σ*=3.2%（对照组）")
    _set_run_font(r, 11, bold=False, color=RGBColor(0x40, 0x40, 0x40))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}"
                  "　·　窗口：IS 2019-05-31 → 2023-12-31，OOS 2024-01-01 → 2025-07-31")
    _set_run_font(r, 10, color=RGBColor(0x60, 0x60, 0x60))

    # ============= §1 背景 =============
    add_heading_zh(doc, "1. 背景与实验设置", level=1)
    add_para(doc,
        "本报告基于已定型的 v6 两层组合（q=0.20 ε=0.30，层1 invvol × lw_erc，"
        "无 trend gate）加装整本波动率目标（vol targeting）杠杆层，"
        "对比在同一底层组合、同一 vol 目标机制下，将杠杆上限 cap 从 2.0 提升至 5.0、"
        "同时将波动率目标 σ* 从 3.2 %（年化）提升至 6.4 % 的效果。"
        "两条主推为 Base 风险预算（55/20/10/15）分别在 GC007 与 DR007 代理融资曲线下的表现；"
        "EW 风险预算（25/25/25/25）作为对照组以说明稳健性。")
    add_para(doc,
        "非-α 板块采用代表性子集（rep-set）。 具体来说，将 6 个非-α 板块（"
        "债券利率、债券信用、跨境DM、跨境HK、金属、其他商品）"
        "通过年度聚类降维压缩至代表性 ETF 子集，将执行头寸数由 53 只降至 31 只（−41 %）。"
        "α 板块（broad_cn、sector_cn）保持 finalist 选券不变。"
        "此代表性子集在两个 cap 的实验中均一致使用，不影响 cap 提升的对比结论。",
        size=10.5)
    add_para(doc,
        "其他所有设置——σ_est 估计器（weekly EWMA-52）、"
        "融资曲线（GC007 或 DR007 代理，即 SHIBOR-1W）、"
        "现金 carry（DR007 代理累计到账内残余现金）、"
        "交易成本（债券 ETF 2 bp/side，其他 10 bp/side）、"
        "调仓节奏（周度 W-FRI）——与原有实验一致，仅 σ* 与 cap 两个参数改变。",
        size=10.5)
    add_para(doc,
        "本报告不改动 v6 生产组合，仅为杠杆上限提升的效果研究。",
        bold=True)

    # ============= §2 主要结果对比 =============
    add_heading_zh(doc, "2. 主要结果对比", level=1)
    add_para(doc,
        "两条黄色高亮行为 Base 风险预算下 cap=5 σ*=6.4 % 的主线结果；"
        "紧邻的非高亮行为对应 cap=2 σ*=3.2 % 的对照。"
        "EW 两行为对照组。")

    add_heading_zh(doc, "2.1 IS + OOS 池化窗口汇总", level=2)
    build_summary_table(doc, "is+oos")

    add_heading_zh(doc, "2.2 IS 窗口（2019-05-31 → 2023-12-31）汇总", level=2)
    build_summary_table(doc, "is")

    add_heading_zh(doc, "2.3 OOS 窗口（2024-01-01 → 2025-07-31）汇总", level=2)
    build_summary_table(doc, "oos")

    add_heading_zh(doc, "2.4 分年度净 CAGR / 最大回撤 / L̄ —— Base RB", level=2)
    add_para(doc,
        "每年两行：第一行显示年内 CAGR（净）与年内平均杠杆倍数 L̄；"
        "第二行显示最大回撤 · 下跌时长（peak → trough 周数）→ 回补时长"
        "（trough → 首次回补前峰的周数，标 “未回补” 者截至 2025-07-31 尚未回补）。",
        size=9.5)
    add_para(doc,
        "关于最大回撤的度量口径：本报告分年度 MaxDD 使用连续净值路径"
        "（cummax 覆盖整段 IS+OOS 序列，非按年 reset 至 1.0）。这样得到的年内 MaxDD "
        "与 §2.1-2.3 汇总表使用的 MaxDD 完全同源，且不会出现年内回撤"
        "机械大于整段窗口回撤的伪失真（后者是常见的 per-year cumsum-from-1 度量偏差）。"
        "2019 与 2025 为部分年份（预热截止 2019-05-31，OOS 截止 2025-07-31），"
        "年化 CAGR 由 compound 方法机械放大，仅作方向参考。",
        size=9.5)
    build_per_year_table(doc)

    add_heading_zh(doc, "2.5 分年度净 CAGR / 最大回撤 / L̄ —— EW RB（对照）", level=2)
    build_per_year_table_ew(doc)

    # ============= §3 图 =============
    add_heading_zh(doc, "3. 图表", level=1)

    add_heading_zh(doc, "3.1 净值曲线", level=2)
    add_figure(doc, "fig_d_nav.png",
               "图1：常数名义 NAV = 1 + Σ 周度 net_ret。"
               "实线为 cap=5 σ*=6.4 %（主线）；虚点线为 cap=2 σ*=3.2 %（对照）。"
               "红色 = GC007 融资；蓝色 = DR007 代理融资。"
               "Base 全不透明，EW 淡化。终末 NAV：Base cap=5 GC007 ≈ 1.61、"
               "cap=2 GC007 ≈ 1.38；Base cap=5 DR007 ≈ 1.62、cap=2 DR007 ≈ 1.38。")

    add_heading_zh(doc, "3.2 回撤曲线", level=2)
    add_figure(doc, "fig_d_drawdown.png",
               "图2：NAV 相对历史峰值的回撤（%）。"
               "cap=5 主线在 2020-03、2022-03、2023-08 等波动集中区间回撤显著加深，"
               "但整体仍在 −6 % 以内；cap=2 对照组同期回撤在 −3 % 以内。"
               "回撤放大幅度大致与平均杠杆倍数成比例——这是设计的直接结果。")

    add_heading_zh(doc, "3.3 杠杆倍数 L_t 时间序列", level=2)
    add_figure(doc, "fig_d_L_path.png",
               "图3：L_t 走势。cap=5 主线在 IS 期（2019–2023）L̄ ≈ 2.1（Base）/ 2.3（EW），"
               "接近 boss 提出的 '正常情况约 2 倍' 目标。"
               "但在 2024–2025 的低波环境下，L_t 单调爬升至 4.5+（Base）和贴顶 5.0（EW），"
               "说明在低波市况中 σ*=6.4 % 需要更多的杠杆才能达到波动目标。"
               "cap=2 对照组则在 2024 年即已贴顶 2.0（虚点线上方水平段）。")

    # ============= §4 久期风险 =============
    add_heading_zh(doc, "4. 久期风险披露", level=1)
    add_para(doc,
        "整本久期使用杠杆后权重（L_t · W_name）乘以静态 KRD 加权计算。"
        "由于 v6 finalist 组合本身债券占比较高（Base 风险预算下利率+信用债 ≈ 30 % NAV，"
        "invvol 加权后由于债券波动低于股票，实际 NAV 占比更高），"
        "整本久期几乎线性随 L̄ 放大——这是本设计新增的主要结构性风险。")

    # duration comparison table
    dur_header = ["组合", "窗口", "久期均值 (年)", "久期 P95 (年)",
                  "100 bp 平行上行冲击（按 P95 久期估算）"]
    dur_rows = []
    for pair_label, c_cell, d_cell in PAIRS:
        s_c = load_summary(c_cell)
        s_d = load_summary(d_cell)
        for w_label, w_key in [("IS+OOS", "is+oos"), ("OOS", "oos")]:
            rc = s_c.loc[w_key]
            rd = s_d.loc[w_key]
            dur_rows.append([
                f"{CELL_LABEL[c_cell]}", w_label,
                f"{rc['book_duration_yr_mean']:.2f}",
                f"{rc['book_duration_yr_p95']:.2f}",
                f"−{rc['book_duration_yr_p95']:.2f} %",
            ])
            dur_rows.append([
                f"{CELL_LABEL[d_cell]}", w_label,
                f"{rd['book_duration_yr_mean']:.2f}",
                f"{rd['book_duration_yr_p95']:.2f}",
                f"−{rd['book_duration_yr_p95']:.2f} %",
            ])
    add_table(doc, dur_header, dur_rows,
              col_widths_cm=[6.0, 1.5, 2.2, 2.2, 4.2])

    add_para(doc,
        "读表要点：",
        bold=True)
    add_bullet(doc,
        "cap=5 主线（Base）IS+OOS 池化久期均值约 10.6 年（Base）/ 12.0 年（EW），"
        "OOS 单窗口均值 13.4 年（Base）/ 15.3 年（EW），"
        "P95 更高——在最长久期的一周里，100 bp 平行利率上行会对净值造成 13–18 % 冲击。")
    add_bullet(doc,
        "cap=2 对照组（Base）IS+OOS 池化久期均值仅约 5.5 年，OOS 单窗口均值 6.5 年，"
        "P95 约 7.4 年——100 bp 冲击约 7 % 净值。")
    add_bullet(doc,
        "cap 提升 2.5 倍带来的久期放大约 2 倍（因为 L̄ 提升不到 2 倍，且债券占比在 EW 下更均衡），"
        "但 100 bp 上行冲击接近线性放大。若市场进入利率上行周期，此风险需在事前明确管控。")
    add_bullet(doc,
        "静态 KRD 表位于 v6/leverage/_common_leverage.py，"
        "近似精度约 ±0.5 年；使用 30 年国债 ETF（511090）= 20 年、"
        "10 年国债/国开（511130/511260/511270）= 8 年等显式覆盖，其余按板块默认。")

    # ============= §5 观察 =============
    add_heading_zh(doc, "5. 主要观察", level=1)
    add_bullet(doc,
        "CAGR 提升显著：Base 风险预算下 IS+OOS 池化 CAGR 由 5.28 %（cap=2）提升至 7.81 %"
        "（cap=5 GC007）、由 5.33 % 提升至 8.07 %（cap=5 DR007 代理），"
        "相当于 +2.5 到 +2.8 pp/年 的绝对回报增益。"
        "OOS 单窗口 CAGR 由 6.6 % 跃升至 11.3 %（GC007）、11.8 %（DR007 代理）。")
    add_bullet(doc,
        "回撤等比放大：Base 池化最大回撤由 −2.55 % 加深至 −5.59 %（GC007）/ −5.44 %（DR007代理），"
        "EW 由 −2.40 % 加深至 −4.41 %。绝对幅度仍在 6 % 以内，"
        "无任何 cell 逼近 10 %。放大幅度大致等比于 L̄ 的放大倍数。")
    add_bullet(doc,
        "回撤时长指标：cap=5 主线最深的 2022 年下跌从 2022-01-14 开始、2022-10-21 见底（下跌 40 周），"
        "然后用 32 周回补到之前的峰值（Base GC007）——即完整水下时长约 72 周（18 个月）。"
        "cap=2 对照在同期下跌 40 周、回补 12 周，水下约 52 周。"
        "水下期间延长约 40 % 是本设计新增的隐性成本，需与 CAGR 提升一并权衡。")
    add_bullet(doc,
        "正常情况下杠杆约 2 倍达成：Base RB 下 IS 期平均 L̄ 为 2.14，EW 为 2.34。"
        "但 OOS 低波环境将 Base L̄ 推高至 3.63，EW 至 4.20（EW 在 18 % 的 OOS 周次贴顶 5.0）。"
        "说明若希望在低波市况下仍保持 L̄ ≈ 2，仅仅提高 σ* 并不够——需要更严格的 L 硬约束或"
        "对 σ* 引入低波调节机制。这是设计的一个可讨论方向。")
    add_bullet(doc,
        "融资成本明显放大：Base GC007 融资成本由 65 bp/年（cap=2）跃升至 335 bp/年（cap=5），"
        "DR007 代理由 58 bp/年 至 297 bp/年，大致等比于 (L̄ − 1) 的放大倍数。"
        "DR007 代理相对 GC007 稳定便宜约 30–40 bp/年，"
        "但如前所述该数值为下界估计，实际执行成本可能更接近 GC007。")
    add_bullet(doc,
        "波动目标机制稳健：L_t 的时序（图3）显示 vol targeter 在 2022 年"
        "股债双杀期自动降至 L ≈ 1–1.5，无过度放大风险；"
        "在 2024–2025 低波期自动提升至 cap 附近，充分利用杠杆额度。"
        "这一自适应机制在 cap 从 2 提升至 5 后依然如此工作。")
    add_bullet(doc,
        "选择建议：若可以接受 OOS 15 年左右的久期敞口与相应的利率上行风险，"
        "推荐 Base RB + cap=5 + σ*=6.4 % 主线；"
        "若倾向更保守的久期敞口，可考虑 cap 保持 2 但 σ* 略调高的中间方案，"
        "或在本框架外引入 L 的硬性上限约束（例如 L ≤ 3）。")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
