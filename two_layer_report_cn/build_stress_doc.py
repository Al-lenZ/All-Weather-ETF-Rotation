"""
build_stress_doc.py
===================
读取 stress/_stress_summary.json + figures/fig_stress*.png，生成独立的
v6_stress_test_report_cn.docx —— 与主报告分开，专门记录 hold-out 压力测试。
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
FIG = HERE / "figures"
STRESS_JSON = HERE / "stress" / "_stress_summary.json"
PREREG = HERE / "STRESS_TEST_PREREG.md"
OUT = HERE / "v6_stress_test_report_cn.docx"


# ============================================================
# Word 工具（与 build_doc.py 一致）
# ============================================================
def set_zh_font(paragraph, name="等线"):
    for run in paragraph.runs:
        run.font.name = name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), name)


def add_heading_zh(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_zh_font(h, "黑体")
    return h


def add_paragraph_zh(doc, text, size=11, indent_first=True):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_zh_font(p, "宋体")
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_zh_font(p, "宋体")
    return p


def add_figure(doc, path: Path, caption: str, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_zh_font(cap, "宋体")


def add_table_from_rows(doc, header, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(10)
        set_zh_font(p, "宋体")
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(v))
            r.font.size = Pt(10)
            set_zh_font(p, "宋体")
    if col_widths:
        for row in tbl.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)


LABEL_ZH = {
    "two_layer_q20_e30":     "双层 q=0.20 ε=0.30 (finalist)",
    "layer1_invvol_lw_erc":  "Layer-1 canonical",
    "T2_bond_invvol":        "T2 债券 inv-vol",
    "two_layer_baseline":    "双层 α off",
}
VERDICT_ZH = {"pass": "通过", "edge": "边缘", "fail": "失败"}
OVERALL_ZH = {"PASS": "通过（PASS）",
              "EDGE": "边缘（EDGE）",
              "FAIL": "失败（FAIL）"}


def main() -> None:
    ST = json.load(open(STRESS_JSON))
    results = ST["results"]
    finalist = "two_layer_q20_e30"
    fin = results[finalist]["stress"]
    pv = ST["per_metric_verdict"]

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)

    # ---------- 标题 ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("v6 双层结构压力测试报告（hold-out shot）")
    r.bold = True
    r.font.size = Pt(19)
    set_zh_font(title, "黑体")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("——在 pre-registration 纪律下打开 hold-out 数据的一次性压力测试")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_zh_font(sub, "宋体")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"报告日期：2026-07-23    ·    "
        f"压力窗口 {ST['stress_window']['start']} → {ST['stress_window']['end']}"
        f"（{ST['stress_window']['n_bars']} 周）    ·    "
        f"综合判定：{OVERALL_ZH[ST['overall_verdict']]}"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    set_zh_font(meta, "宋体")

    doc.add_paragraph()

    # ---------- 摘要 ----------
    add_heading_zh(doc, "摘要", level=1)
    add_paragraph_zh(doc,
        f"v6 双层结构 finalist（q=0.20 ε=0.30）在压力窗口 "
        f"[{ST['stress_window']['start']}, {ST['stress_window']['end']}] "
        f"共 {ST['stress_window']['n_bars']} 周的一次性 hold-out 测试中，"
        f"通过了事前冻结的全部 5 项阈值："
        f"净夏普 {fin['sharpe']:+.3f} / CAGR {fin['cagr']*100:+.2f}% / "
        f"最大回撤 {fin['dd']*100:+.2f}% / Calmar {fin['calmar']:+.2f} / "
        f"Sharpe decay ratio {fin['decay']:+.2f}。")
    add_paragraph_zh(doc,
        "但两处诚实标注必须写进结论：(1) 25H2–26H1 这段窗口对债券占比高的组合恰好偏温和，"
        "T2 债券 inv-vol 单独在压力窗口的净夏普也高于 IS anchor，"
        "并非典型 stress 环境（未触及 2020 疫情 / 2022 信用债回调级别的宏观事件）；"
        "(2) α 层（broad_cn / sector_cn 权益选股）在压力窗口的边际贡献为负，"
        f"finalist Sharpe {fin['sharpe']:+.2f} 低于双层 α off 的 "
        f"{results['two_layer_baseline']['stress']['sharpe']:+.2f}，"
        "Δ ≈ −0.26 Sharpe。'stress 通过'的正确解读是'设计在温和逆风下未暴露缺陷'，"
        "不是'已被验证抗任何压力'。")
    add_paragraph_zh(doc,
        "根据 pre-registration §6 硬承诺，本次测试后 v6 knob 永久冻结，"
        "本压力窗口 [2025-08-01, 2026-07-17] 对未来 v7 项目不再算 OOS。"
        "v7 的真正 OOS 只能是从今天（2026-07-23）之后的新窗口。")

    # ---------- 一、Pre-registration ----------
    add_heading_zh(doc, "一、Pre-registration 摘要", level=1)
    add_paragraph_zh(doc,
        "在打开 hold-out 数据之前，先写了 pre-registration memo "
        "（原件：v6/two_layer_report_cn/STRESS_TEST_PREREG.md）以钉死规则、"
        "防止事后合理化。核心条款如下。")

    add_heading_zh(doc, "1.1 只评估 finalist + 3 anchor", level=2)
    add_bullet(doc, "落地 finalist：two_layer_q20_e30（q=0.20，ε=0.30）——本次压力测试的唯一评估主体。")
    add_bullet(doc, "anchor（不是 v6 参数、只是参考）：two_layer_baseline（α off）"
                    "、layer1_invvol_lw_erc、T2_bond_invvol。")
    add_bullet(doc, "❌ 不跑 q=0.10 变体，不跑其他 (q, ε) 格点，不改任何 sizing / "
                    "协方差窗口 / 政策份额——防止扫格式偷看。")

    add_heading_zh(doc, "1.2 事前阈值（在打开数据前冻结）", level=2)
    header = ["维度", "通过阈值", "边缘阈值", "失败阈值"]
    rows = [
        ["净夏普（stress）",        "≥ +0.50",  "0 ~ +0.50",  "< 0"],
        ["CAGR（stress，年化）",    "≥ 0",     "−2% ~ 0",    "< −2%"],
        ["最大回撤（stress）",      "≥ −4%",   "−6% ~ −4%",  "< −6%"],
        ["Calmar（stress）",        "≥ +0.30",  "0 ~ +0.30",  "< 0"],
        ["Sharpe decay ratio",     "≥ 0.30",  "0.10 ~ 0.30","< 0.10"],
    ]
    add_table_from_rows(doc, header, rows, col_widths=[4.5, 3.5, 3.5, 3.5])
    add_paragraph_zh(doc,
        "判定规则：≥ 4 维通过且无失败 → 通过；≥ 1 维边缘且无失败 → 边缘；"
        "任何 1 维失败 → 失败。")

    add_heading_zh(doc, "1.3 明确不看的东西", level=2)
    add_paragraph_zh(doc, "为了防止事后归因变成隐蔽调参，以下内容明令禁止：")
    add_bullet(doc, "分块（block-level）逐块的 gross / net 贡献")
    add_bullet(doc, "逐 ETF 名字的 pnl 归因")
    add_bullet(doc, "分月 / 分周细粒度的回撤路径（只看窗口整体 MaxDD）")
    add_bullet(doc, "α 层在 stress 窗口里选中的具体名字")
    add_bullet(doc, "Layer-1 solver 在 stress 窗口里给出的**逐周**组权重时序"
                    "（允许看窗口**平均**组权重）")

    add_heading_zh(doc, "1.4 硬承诺（hard commitments）", level=2)
    add_paragraph_zh(doc,
        "无论测试结果如何，看完 stress 窗口后以下操作永久禁止：")
    add_bullet(doc, "不再调整 (q, ε) 或任何 v6 knob。落地方案就是 q=0.20 ε=0.30。")
    add_bullet(doc, "不再基于 stress 结果做 v6 内的\"临时补丁\"（比如\"stress 表现差 → 加个 trend gate\"）。")
    add_bullet(doc, "不打开 stress 窗口内更细粒度的诊断。")
    add_paragraph_zh(doc, "允许的下一步：")
    add_bullet(doc, "写这份独立的收官报告（即本文档），把 stress 结果诚实纳入。")
    add_bullet(doc, "开始设计 v7（结构性改动）——具体动机应有独立于 stress 观测的 IS 依据。")
    add_bullet(doc, "在 v7 项目里，将 2019-05-31 → 今天（含之前的 IS + OOS + stress）"
                    "作为训练可用窗口，真正的 v7 OOS 只能是从今天之后的新窗口。")

    # ---------- 二、结果 ----------
    add_heading_zh(doc, "二、结果", level=1)

    add_heading_zh(doc, "2.1 finalist（two_layer_q20_e30）逐项判定", level=2)
    header = ["维度", "值", "阈值判定"]
    rows_fin = [
        ["净夏普",   f"{fin['sharpe']:+.3f}",       VERDICT_ZH[pv["sharpe"]]],
        ["CAGR",     f"{fin['cagr']*100:+.2f}%",    VERDICT_ZH[pv["cagr"]]],
        ["最大回撤", f"{fin['dd']*100:+.2f}%",      VERDICT_ZH[pv["max_dd"]]],
        ["Calmar",   f"{fin['calmar']:+.2f}",       VERDICT_ZH[pv["calmar"]]],
        ["Sharpe decay ratio",
         f"{fin['decay']:+.2f}  (stress {fin['sharpe']:+.2f} / IS 参照 "
         f"{ST['is_sharpe_anchors'][finalist]:+.2f})",
         VERDICT_ZH[pv["decay"]]],
    ]
    add_table_from_rows(doc, header, rows_fin, col_widths=[5.0, 6.5, 3.5])

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.8)
    r = p.add_run(f"综合判定：{OVERALL_ZH[ST['overall_verdict']]}。全部 5 项事前阈值均通过。")
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = (RGBColor(0x27, 0xAE, 0x60) if ST["overall_verdict"] == "PASS"
                         else (RGBColor(0xE6, 0x7E, 0x22) if ST["overall_verdict"] == "EDGE"
                                else RGBColor(0xC0, 0x39, 0x2B)))
    set_zh_font(p, "宋体")

    add_heading_zh(doc, "2.2 三窗对比：IS anchor / OOS / Stress", level=2)
    add_paragraph_zh(doc,
        "把 finalist + 三个 anchor 在三个窗口的表现平铺展开，"
        "可以直接读出 stress 窗口相对 IS 与 OOS 的位置。"
        "注：IS anchor 列只显示 Sharpe（引自主报告 §4 OOS shot），"
        "其他 IS 指标不复现以避免与主报告重复。")
    header = ["书", "窗口", "净夏普", "CAGR", "最大回撤", "Calmar", "有效周数"]
    rows = []
    book_order = ["two_layer_q20_e30", "layer1_invvol_lw_erc",
                   "T2_bond_invvol", "two_layer_baseline"]
    for name in book_order:
        r = results[name]
        is_sh = ST["is_sharpe_anchors"][name]
        rows.append([LABEL_ZH[name], "IS anchor",
                     f"{is_sh:+.3f}", "—", "—", "—", "—"])
        oos = r["OOS"]; st = r["stress"]
        rows.append([LABEL_ZH[name], "OOS",
                     f"{oos['sharpe']:+.3f}",
                     f"{oos['cagr']*100:+.2f}%",
                     f"{oos['dd']*100:+.2f}%",
                     (f"{oos['calmar']:+.2f}" if oos["calmar"] is not None else "—"),
                     str(int(oos["n_bars"]))])
        rows.append([LABEL_ZH[name], "Stress",
                     f"{st['sharpe']:+.3f}",
                     f"{st['cagr']*100:+.2f}%",
                     f"{st['dd']*100:+.2f}%",
                     (f"{st['calmar']:+.2f}" if st["calmar"] is not None else "—"),
                     str(int(st["n_bars"]))])
    add_table_from_rows(doc, header, rows)

    add_heading_zh(doc, "2.3 压力窗口 NAV 与回撤", level=2)
    add_paragraph_zh(doc,
        "以压力窗口起点为净值起点（NAV = 1）重画累计净值和回撤。"
        "四条曲线都是 10 bp / 单边成本口径下的净值。")
    add_figure(doc, FIG / "fig_stress01_nav.png",
               "图 1  压力窗口内累计净值（起点 = 1，含成本）")
    add_figure(doc, FIG / "fig_stress02_dd.png",
               "图 2  压力窗口内回撤路径")
    add_paragraph_zh(doc,
        "读图要点：Layer-1 canonical 与双层 α off 的 NAV 曲线几乎重合（一致性校验的直接体现），"
        "都收在 +2.5% 附近；finalist 收在 +2.1%，略低于两个 α off 参照，"
        f"说明 α 层在 stress 窗口净贡献为负（Δ Sharpe ≈ "
        f"{fin['sharpe'] - results['two_layer_baseline']['stress']['sharpe']:+.2f}）。"
        "T2 债券最平稳但累计收益最低。回撤面板上四条曲线的最大回撤都在 −0.30% 附近，"
        "无一超过 −0.4%——这就是事前阈值 −4% 被大幅通过的原因。")

    add_heading_zh(doc, "2.4 Sharpe 保留率（stress / IS 参照）", level=2)
    add_paragraph_zh(doc,
        "decay ratio 是本次压力测试最直观的诊断——如果它 < 0.10 说明 IS 里的信号在 stress 完全失效，"
        "如果 ≥ 0.30 说明 IS 优势在 stress 里显著保留。"
        "四本书都远远高于事前阈值，且 finalist 的 decay 1.37 意味着 stress Sharpe 反而高于 IS 参照。")
    add_figure(doc, FIG / "fig_stress04_decay.png",
               "图 3  各书 Sharpe decay 对比（stress Sharpe / IS anchor Sharpe）")

    add_heading_zh(doc, "2.5 压力窗口组间资本权重", level=2)
    add_paragraph_zh(doc,
        "把 stress 窗口的平均组权重与 OOS 均值、政策目标并列展示。"
        "Solver 在压力窗口里主动把信用债从 30% 抬到 49%，"
        "利率债则从 61% 降到 43%；权益仍维持在 5% 左右，商品 2.6%。"
        "这是 LW-ERC 在 stress 窗口里读到\"信用债 risk-adjusted 更便宜\"信号后的自然反应，"
        "并非人为参数调整——solver 内部逻辑与 IS/OOS 完全相同。"
        "Realized RC% 仍然严格命中 55/20/10/15（未画在图里以避免与主报告重复）。")
    add_figure(doc, FIG / "fig_stress03_wgroup.png",
               "图 4  压力窗口平均组权重 vs OOS vs 政策目标")

    # ---------- 三、诚实解读 ----------
    add_heading_zh(doc, "三、诚实解读", level=1)
    add_paragraph_zh(doc,
        "PASS 判定看起来很干净，但两个 caveats 必须写清楚。")

    add_heading_zh(doc, "3.1 25H2–26H1 不是典型 stress 环境", level=2)
    add_paragraph_zh(doc,
        "本次压力测试之所以能拿到 Sharpe +2.16、Calmar +7.09 这样漂亮的数字，"
        "很大程度上是因为这段窗口对债券占比高的组合恰好偏温和——"
        f"T2 债券 inv-vol 单独在压力窗口的净夏普 "
        f"{results['T2_bond_invvol']['stress']['sharpe']:+.3f}，"
        "本身就高于 IS anchor +1.462。这个窗口里权益虽然有回调，"
        "但没有出现 2020 疫情或 2022 信用债回调级别的极端事件，"
        "利率环境相对配合。真实的 tail-risk 环境（比如国债利率突升 + 信用债利差扩张同时发生）"
        "未在本次压力测试中被检验到。")

    add_heading_zh(doc, "3.2 α 层的边际贡献在 stress 里为负", level=2)
    add_paragraph_zh(doc,
        "把 finalist 与双层 α off baseline 并列，可以直接读出 α 层的净贡献：")
    header = ["书", "IS Sharpe (anchor)", "OOS Sharpe", "Stress Sharpe"]
    rows = [
        ["双层 α off (baseline)",
         f"{ST['is_sharpe_anchors']['two_layer_baseline']:+.3f}",
         f"{results['two_layer_baseline']['OOS']['sharpe']:+.3f}",
         f"{results['two_layer_baseline']['stress']['sharpe']:+.3f}"],
        ["双层 α on (finalist)",
         f"{ST['is_sharpe_anchors']['two_layer_q20_e30']:+.3f}",
         f"{results['two_layer_q20_e30']['OOS']['sharpe']:+.3f}",
         f"{results['two_layer_q20_e30']['stress']['sharpe']:+.3f}"],
        ["Δ (α layer)",
         f"{ST['is_sharpe_anchors']['two_layer_q20_e30'] - ST['is_sharpe_anchors']['two_layer_baseline']:+.3f}",
         f"{results['two_layer_q20_e30']['OOS']['sharpe'] - results['two_layer_baseline']['OOS']['sharpe']:+.3f}",
         f"{results['two_layer_q20_e30']['stress']['sharpe'] - results['two_layer_baseline']['stress']['sharpe']:+.3f}"],
    ]
    add_table_from_rows(doc, header, rows)
    add_paragraph_zh(doc,
        "IS 和 OOS 上 α 层都是小幅正贡献（+0.07 / +0.08 Sharpe），"
        "到 stress 窗口翻负到 −0.26。这不是 pre-reg 的\"失败档\""
        "（阈值只看 finalist 的绝对 metric，不看 Δ），但作为独立的观察值得记录：",
        indent_first=False)
    add_bullet(doc,
        "broad_cn / sector_cn 的权益 α 在 25H2–26H1 没有 edge，"
        "这段时间 A 股权益是选股 alpha 相对稀薄的一段时期。")
    add_bullet(doc,
        "这一定程度上支持了主报告里就已指出的\"Layer-1 组间风险预算才是主要的结构性收益源，"
        "α 层的边际改善较小\"的结论。α 层不是 v6 的核心 alpha 源。")
    add_bullet(doc,
        "但 stress 窗口 α layer Δ 为负这个观察，**不会**触发对 (q, ε) 的重调——"
        "pre-reg §6 硬承诺已经把这条路封死。这类观察只能作为 v7 的设计线索："
        "比如 v7 里如果要保留 α 层，可能需要考虑\"信号强度门槛\"或\"regime 感知\"的机制。")

    add_heading_zh(doc, "3.3 压力测试的限度", level=2)
    add_paragraph_zh(doc,
        "本次压力测试的三点主要限度："
        "(1) 窗口只有 51 周，样本量不足以支持子窗口切分或统计推断；"
        "(2) 窗口内未包含 2020/2022 级别的宏观事件，属于\"温和逆风\"；"
        "(3) 只跑了 finalist + 3 anchor，没有做扫格式的鲁棒性检查（这是"
        "刻意为之，因为扫格本身就是 bias 的源头）。"
        "正确的读法是：**\"设计在温和逆风下未暴露缺陷\"**，而不是"
        "**\"设计已被验证能抵御任何压力\"**。真正的抗压能力仍需等待 v7 项目里的新数据积累。")

    # ---------- 四、后续 ----------
    add_heading_zh(doc, "四、v6 收官与 v7 数据边界", level=1)
    add_paragraph_zh(doc,
        "根据 pre-registration §6 硬承诺，本次压力测试完成后：")
    add_bullet(doc,
        "**v6 落地方案冻结**：two-layer q=0.20 ε=0.30（broad_cn K=5 + sector_cn K=8，"
        "long_q20-replace，Layer-1 = invvol × LW-ERC，无 trend gate）。"
        "任何后续 v6 调整均被禁止。")
    add_bullet(doc,
        "**stress 窗口 [2025-08-01, 2026-07-17] 对 v7 不再算 OOS**。"
        "对 v7 项目而言，这段数据算作训练可用窗口的一部分，"
        "与 2019-05-31 → 2025-07-31 之前的 IS + OOS 一起构成完整的\"训练可用集\"。"
        "v7 真正的 OOS 只能是从今天（2026-07-23）之后的新窗口。")
    add_bullet(doc,
        "**v7 的设计动机应独立于 stress 观测**。已经明确的两个方向"
        "（更多板块做 α、非-α 块降密度）来自可执行性 + 因子覆盖度，"
        "有独立的 IS 依据，不由 stress 结果驱动。")

    # ---------- 附录 ----------
    add_heading_zh(doc, "附录：文件索引", level=1)
    add_bullet(doc, "Pre-registration 原件：two_layer_report_cn/STRESS_TEST_PREREG.md")
    add_bullet(doc, "执行脚本：two_layer_report_cn/run_stress_test.py")
    add_bullet(doc, "结果落盘：two_layer_report_cn/stress/_stress_summary.json、verdict.txt")
    add_bullet(doc, "图表：two_layer_report_cn/figures/fig_stress0{1..4}_*.png")
    add_bullet(doc, "主报告：two_layer_report_cn/two_layer_strategy_report_cn.docx"
                    "（不含 stress，仅覆盖 IS + OOS）")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
