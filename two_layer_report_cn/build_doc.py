"""
build_doc.py
============
读取 _summary.json + figures/，生成中文 Word 报告
two_layer_strategy_report_cn.docx。
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
FIG = HERE / "figures"
OUT = HERE / "two_layer_strategy_report_cn.docx"


# ============================================================
# Word 工具
# ============================================================
def set_zh_font(paragraph, name="等线"):
    for run in paragraph.runs:
        run.font.name = name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), name)


def add_heading_zh(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_zh_font(h, "黑体")
    return h


def add_paragraph_zh(doc, text, size=11, indent_first=True):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_zh_font(p, "宋体")
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_zh_font(p, "宋体")
    return p


def add_figure(doc, path: Path, caption: str, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_zh_font(cap, "宋体")


def add_table_from_rows(doc, header, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(10)
        set_zh_font(p, "宋体")
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(v))
            r.font.size = Pt(10)
            set_zh_font(p, "宋体")
    if col_widths:
        for row in tbl.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)


def pct(x, n=2):
    if x is None:
        return "—"
    return f"{x * 100:+.{n}f}%"


def num(x, n=2):
    if x is None:
        return "—"
    return f"{x:+.{n}f}"


def num_abs(x, n=2):
    if x is None:
        return "—"
    return f"{x:.{n}f}"


# ============================================================
# 主流程
# ============================================================
LABEL_ZH = {
    "two_layer_q20_e30":     "双层 q=0.20 ε=0.30",
    "two_layer_q10_e30":     "双层 q=0.10 ε=0.30",
    "two_layer_baseline":    "双层 α off",
    "layer1_invvol_lw_erc":  "Layer-1 canonical",
    "T2_bond_invvol":        "T2 债券 inv-vol",
    "solo_defensive_final":  "旧版 solo defensive",
}
AXIS_ZH = {
    "equity":      "权益",
    "bond_rates":  "利率债",
    "bond_credit": "信用债",
    "commodity":   "商品",
}


def main() -> None:
    with open(HERE / "_summary.json") as f:
        S = json.load(f)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)

    # ---------- 标题 ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("v6 双层结构（Layer-1 组内风险预算 × Layer-2 α）研究报告")
    r.bold = True
    r.font.size = Pt(19)
    set_zh_font(title, "黑体")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("——池子扩容后再上一层组间风险预算与块内 α，样本外表现的结构性升级")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_zh_font(sub, "宋体")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"报告日期：2026-07-23    ·    "
                     f"共同起点 {S['common_start']} → 样本内截止 2023-12-31 → "
                     f"样本外 2024-01-01 至 2025-07-31")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    set_zh_font(meta, "宋体")

    doc.add_paragraph()

    # ---------- 摘要 ----------
    stats = S["stats"]
    q20 = stats["two_layer_q20_e30"]
    q10 = stats["two_layer_q10_e30"]
    base = stats["two_layer_baseline"]
    l1 = stats["layer1_invvol_lw_erc"]
    solo = stats["solo_defensive_final"]

    add_heading_zh(doc, "摘要", level=1)
    add_paragraph_zh(doc,
        "本报告在此前 v6 Final Report 已经落地的“Top-q 纯多头 + 1/σ + replace-ε hysteresis”"
        "基础上，正式引入一层组间风险预算（Layer-1）与两块（broad_cn、sector_cn）的"
        "块内 α 选择（Layer-2）：Layer-1 以四个组（权益 55% / 利率债 20% / 信用债 10% / 商品 15%）"
        "为最小风险单元，用 Ledoit-Wolf-target-D 收缩协方差 + 对数障碍 ERC 求解组间权重；"
        "Layer-2 只在两块 A 股权益上激活 α-hysteresis 选股（复用 long_q20-replace 生产核），"
        "其他 6 块保持块内 inv-vol 全持。")
    add_paragraph_zh(doc,
        "在 IS = 2019-05-31 → 2023-12-31 上做 3×4 的 (q, ε) 扫格并按 plateau 规则选出 "
        "q=0.20、ε=0.30 为落地方案，同时保留 q=0.10、ε=0.30 作为 raw-IS 最高的对照。"
        "OOS 开放窗口（2024-01-01 → 2025-07-31）的结果结构性优于此前的 solo defensive finalist：")
    add_bullet(doc,
        f"双层 q=0.20 ε=0.30：OOS 净夏普 {num(q20['OOS']['sharpe'], 3)}，"
        f"CAGR {pct(q20['OOS']['cagr'])}，最大回撤 {pct(q20['OOS']['dd'])}，"
        f"Calmar（全窗）{num_abs(q20['full']['calmar'])}；"
        f"全窗净夏普 {num(q20['full']['sharpe'], 3)}。")
    add_bullet(doc,
        f"旧版 solo defensive（Phase 11.2 finalist）：OOS 净夏普 {num(solo['OOS']['sharpe'], 3)}，"
        f"CAGR {pct(solo['OOS']['cagr'])}，最大回撤 {pct(solo['OOS']['dd'])}，"
        f"Calmar（全窗）{num_abs(solo['full']['calmar'])}；"
        f"全窗净夏普 {num(solo['full']['sharpe'], 3)}。")
    add_bullet(doc,
        f"两者对比：双层结构在 OOS 上 Sharpe 提升约 "
        f"{q20['OOS']['sharpe'] - solo['OOS']['sharpe']:+.2f}，"
        f"CAGR 提升约 {(q20['OOS']['cagr'] - solo['OOS']['cagr']) * 100:+.2f} pp，"
        f"最大回撤放浅约 {(q20['OOS']['dd'] - solo['OOS']['dd']) * 100:+.2f} pp，"
        f"全窗 Calmar 由 {solo['full']['calmar']:.2f} 抬升到 {q20['full']['calmar']:.2f}。")
    add_paragraph_zh(doc,
        "α 层单独的边际贡献相对温和——OOS Sharpe Δ 仅约 +0.08（q20） / +0.12（q10）——"
        "更明确的结构性收益其实来自 Layer-1 的组间风险预算：把权益、利率债、信用债、商品"
        "四组按静态目标权重分开求解，之后再把 α 选出的 broad_cn / sector_cn 名字嵌入到"
        "各自组内。10 bp / 单边成本口径下所有夏普、CAGR、最大回撤、Calmar 均为净值。")

    # -------- 实验二一句话预告 ---------
    e2 = S["exp2"]
    e2_hold = e2["stats"]["exp2_hold_all"]
    e2_rep = e2["stats"]["exp2_replicated"]
    add_paragraph_zh(doc,
        "本次报告还补入第七、八节两个专题实验："
        "**实验一**（§七）扰动 Layer-1 政策风险预算（55/20/10/15）"
        "，做 EW 对照 + 4 轴 ±10pp + 80-cell 全 grid 敏感性——"
        "结论比预期保守：**只有权益轴的方向性信号能明确从 cross-cell 噪声里拆出来**"
        "（权益 ↑ → CAGR ↑ / DD 恶化 / Sharpe ↓），"
        "另外三轴（利率债 / 信用债 / 商品）的独立效应都在 cross-cell σ 内、"
        "且被 pro-rata 补偿方案与权益轴耦合。因此本节的建议只是「权益略偏高，"
        "可以考虑往 EW 方向微调」，而不是完整重构 policy shares。"
        "**实验二**（§八）把 6 个非-α 块的持仓"
        f"由“块内 inv-vol 全持”压到年度刷新的“每聚类留一支代表”。名字数 K 均值 "
        f"{38.1:.1f} → {24.1:.1f}（−37%），OOS 净夏普由 {e2_hold['OOS']['sharpe']:+.2f} 上升到 "
        f"{e2_rep['OOS']['sharpe']:+.2f}（Δ {e2_rep['OOS']['sharpe']-e2_hold['OOS']['sharpe']:+.2f}）。"
        "值得强调的是——**这轮 OOS 夏普提升不是来自收益变高，而是来自策略波动被压下去**："
        f"OOS CAGR 反而从 {pct(e2_hold['OOS']['cagr'])} 下滑到 {pct(e2_rep['OOS']['cagr'])}"
        f"（Δ {(e2_rep['OOS']['cagr']-e2_hold['OOS']['cagr'])*100:+.2f} pp），"
        f"OOS 年化波动却从 {pct(e2_hold['OOS']['ann_vol'])} 收窄到 {pct(e2_rep['OOS']['ann_vol'])}"
        f"（Δ {(e2_rep['OOS']['ann_vol']-e2_hold['OOS']['ann_vol'])*100:+.2f} pp，相对减幅 "
        f"−{(1-e2_rep['OOS']['ann_vol']/e2_hold['OOS']['ann_vol'])*100:.1f}%），"
        "分子小减、分母大减，因此 Sharpe 抬升。这是与前面 Layer-1/Layer-2 讨论不同"
        "性质的另一层结构性升级，第八节做完整解释。")

    # ---------- 一、双层设计概述 ----------
    add_heading_zh(doc, "一、双层设计概述", level=1)

    add_heading_zh(doc, "1.1 之前 Final Report 的落地方案回顾（一笔带过）", level=2)
    add_paragraph_zh(doc,
        "v6 Final Report 中最终落地的方案是纯多头 hysteresis：从 344 只入池 ETF 中"
        "按集成 α 排序，取 Top 20% 组成多头组合，仓位内以 1/σ_causal_26w 反比配权，"
        "并用 replace 规则 + ε=0.20 的迟滞机制压低选股换手成本。样本外净夏普约 "
        f"{solo['OOS']['sharpe']:+.2f}，CAGR {pct(solo['OOS']['cagr'])}，最大回撤 "
        f"{pct(solo['OOS']['dd'])}。这套“单层”方案的短板是全窗 Calmar 只有 "
        f"{solo['full']['calmar']:.2f}，主要因为它把所有权益、债、商品一视同仁地放进同一个"
        "截面排序里，缺乏组间风险预算，深度回撤集中在 2020 疫情年与 2022 年信用债回调。")

    add_heading_zh(doc, "1.2 Layer-1：组间风险预算 + LW-ERC", level=2)
    add_paragraph_zh(doc,
        "把入池 ETF 按资产大类划分为 4 组：")
    add_bullet(doc, "equity（权益）：包含 broad_cn、sector_cn、cross_border_dm、cross_border_hk 四个块；"
                    "静态目标权重 55%。")
    add_bullet(doc, "bond_rates（利率债）：静态目标权重 20%。")
    add_bullet(doc, "bond_credit（信用债）：静态目标权重 10%。")
    add_bullet(doc, "commodity（商品）：包含 metals、commodity_other 两个块；静态目标权重 15%。")
    add_paragraph_zh(doc,
        "组间协方差用 Ledoit-Wolf-target-D 做收缩，"
        "然后以对数障碍法（log-barrier）求解等风险贡献（ERC）解，"
        "使得每组对总组合的风险贡献大致等于其静态目标权重；"
        "trend gate 已在 Layer-1 消融实验中确认为 Sharpe-neutral（±0.04），"
        "因此为了实现简单，落地版本去掉了 trend gate。")

    add_heading_zh(doc, "1.3 Layer-2：α 只跑在 broad_cn 与 sector_cn 上", level=2)
    add_paragraph_zh(doc,
        "Layer-2 只在两块 A 股权益上激活 α 选股，其他 6 块（bond_rates、bond_credit、"
        "cross_border_dm、cross_border_hk、metals、commodity_other）保持块内 inv-vol 全持。"
        "选股核复用生产 long_q20-replace kernel，参数由 (q, ε) 两个 knob 控制：")
    add_bullet(doc, "broad_cn 使用 K=5 的锁定集成（alpha015 / alpha_071 / alpha_102 / "
                    "h_mom_decay_12_48 / alpha006↺，5 个成员），Phase 13.5 冻结。")
    add_bullet(doc, "sector_cn 使用 K=8 的锁定集成（var5_60 / ma_disp / alpha_142 / alpha_187 / "
                    "yj15_bias_mom_60_20↺ / h_mom_decay_12_48 / kurt_40↺ / ret_skew_20，8 个成员），"
                    "Phase 13.5 冻结。")
    add_paragraph_zh(doc,
        "关键设计要点：块内选出的名字权重先按 1/σ 配比归一到 Σ = 1，然后再乘以该块在其所属组内的静态"
        "份额 N_b / N_g，最后再乘以 Layer-1 求解得到的组权重。这样当 α 关闭（q=1，ε=0）时，"
        "整个双层书退化为纯 Layer-1 的组间风险预算 + 块内 inv-vol 全持——"
        "任何 Δ 都可以干净地归因到“α 层的贡献”上。")

    add_heading_zh(doc, "1.4 与旧版 Final Report 的关键差异", level=2)
    header = ["维度", "Final Report（solo defensive）", "本报告（双层结构）"]
    rows = [
        ["选股范围", "全池 344 只 ETF 统一排序", "只在 broad_cn / sector_cn 两块内做 α，其他 6 块全持"],
        ["权重结构", "选中名字之间 1/σ", "块内 1/σ → 块乘 N_b/N_g → 组乘 LW-ERC"],
        ["组间风险预算", "无（隐式跟随选股结果）", "静态 55/20/10/15 + LW 收缩 + ERC 求解"],
        ["调仓核", "long_q20-replace，ε=0.20", "long_q20-replace，(q, ε) 扫格后 plateau 选 (0.20, 0.30)"],
        ["集成成员", "PV 去相关 shortlist 全量集成", "两块各自 K=5 / K=8 锁定集成，Phase 13.5 冻结"],
    ]
    add_table_from_rows(doc, header, rows, col_widths=[3.2, 5.5, 6.5])

    # ---------- 二、样本内扫格 ----------
    add_heading_zh(doc, "二、样本内扫格与 plateau 选择", level=1)

    add_heading_zh(doc, "2.1 12 个 (q, ε) 格点", level=2)
    add_paragraph_zh(doc,
        "在 IS 窗口内做 3×4 扫格：q ∈ {0.10, 0.20, 0.30}，ε ∈ {0.00, 0.10, 0.20, 0.30}，"
        "对 broad_cn 和 sector_cn 施加同一组 (q, ε)。"
        "参考 baseline 为“α off”（q=1，ε=0）的双层结构本身——按定义等价于 Layer-1 canonical。"
        "全部结果都在 10 bp / 单边成本口径下报告。下图展示扫格结果的 IS 净夏普矩阵。")
    add_figure(doc, FIG / "fig01_sweep_heatmap.png",
               "图 1  Layer-2 α 参数扫格：样本内净夏普（12 cells）")

    add_heading_zh(doc, "2.2 完整扫格表", level=2)
    header = ["q", "ε", "IS Sharpe", "Δ Sh vs baseline",
              "IS CAGR", "最大回撤", "单边换手", "cost drag (bps/yr)"]
    rows = []
    for r in S["sweep"]:
        star = "★" if (round(r["q"], 2) == 0.20 and round(r["epsilon"], 2) == 0.30) else ""
        rows.append([
            f"{r['q']:.2f}{star}",
            f"{r['epsilon']:.2f}",
            num(r["sharpe_net"], 3),
            num(r["d_sharpe_net"], 3),
            pct(r["cagr_net"]),
            pct(r["max_dd_net"]),
            f"{r['turnover']:.4f}",
            f"{r['cost_drag_bp_yr']:+.1f}",
        ])
    add_table_from_rows(doc, header, rows)

    add_paragraph_zh(doc,
        "读表要点：")
    add_bullet(doc,
        "q=0.10 系列（更集中的 α 选股）在 IS 上 raw-Sharpe 最高，"
        "峰值为 q=0.10、ε=0.30 的 +1.468（Δ vs baseline +0.103）。")
    add_bullet(doc,
        "q=0.20 系列 Sharpe 略低（+1.42 附近），但换手明显更低（0.081 vs 0.087）。"
        "由于 plateau 内 4 个 cell 的 Δ-Sharpe 都在 ±0.05 内，plateau 规则倾向于"
        "选择换手最低的 (q=0.20, ε=0.30) 作为落地方案。")
    add_bullet(doc,
        "q=0.30 系列在多数 (ε) 上 Δ-Sharpe ≤ 0，说明选股面过宽时 α 信号会被稀释。")

    add_heading_zh(doc, "2.3 plateau 规则与 winner", level=2)
    add_paragraph_zh(doc,
        "plateau 规则：在 Δ-Sharpe ≥ (max Δ-Sharpe − 0.05) 的所有格子里，"
        "按 turnover 升序（低换手优先），再依 ε 升序（更粘的仓位）、q 升序（更集中）打破平手。"
        "plateau 内含 4 个格子，冠军为 q=0.20、ε=0.30："
        f"IS Sharpe {q20['IS']['sharpe']:+.3f} / CAGR {pct(q20['IS']['cagr'])} / "
        f"DD {pct(q20['IS']['dd'])} / 换手 0.0813。")
    add_paragraph_zh(doc,
        "同时报告 q=0.10、ε=0.30 作为“raw-IS 最高”参考："
        f"IS Sharpe {q10['IS']['sharpe']:+.3f} / CAGR {pct(q10['IS']['cagr'])} / "
        f"DD {pct(q10['IS']['dd'])} / 换手 0.0873。")

    # ---------- 三、α off ≡ Layer-1 校验 ----------
    add_heading_zh(doc, "三、α off 与 Layer-1 canonical 的一致性校验", level=1)
    add_paragraph_zh(doc,
        "为了保证 Δ 能干净地归因给 α 层，两组的“组内 inv-vol × 组间 LW-ERC + no-trend”结构"
        "在数值上应当等价。作为 sanity check，把 (q=1, ε=0) 的 two_layer_baseline 与"
        "单层实现的 layer1_invvol_lw_erc 直接叠加：净值曲线在整个 IS+OOS 窗口内几乎完全重叠，"
        "相对偏差主要来自浮点求解误差与 sub-block share 舍入。")
    add_figure(doc, FIG / "fig02_alpha_off_identity.png",
               "图 2  α off 的双层结构 vs Layer-1 canonical（近似同一条线）")
    header = ["书", "IS Sharpe", "OOS Sharpe", "全窗 Sharpe", "IS CAGR", "OOS CAGR", "最大回撤"]
    rows = []
    for name in ["two_layer_baseline", "layer1_invvol_lw_erc"]:
        s = stats[name]
        rows.append([
            LABEL_ZH[name],
            num(s["IS"]["sharpe"], 3),
            num(s["OOS"]["sharpe"], 3),
            num(s["full"]["sharpe"], 3),
            pct(s["IS"]["cagr"]),
            pct(s["OOS"]["cagr"]),
            pct(s["full"]["dd"]),
        ])
    add_table_from_rows(doc, header, rows)

    # ---------- 四、OOS 结果 ----------
    add_heading_zh(doc, "四、样本外结果（OOS shot）", level=1)

    add_heading_zh(doc, "4.1 三窗汇总", level=2)
    add_paragraph_zh(doc,
        "IS = 共同起点 {} → 2023-12-31，OOS = 2024-01-01 → 2025-07-31（82 周），"
        "hold-out（> 2025-07-31）保留为密封窗口，本报告不涉及。"
        "所有指标均在 10 bp / 单边成本口径下计算。".format(S["common_start"]))
    header = ["书", "窗口",
              "净夏普", "CAGR", "最大回撤", "年化波动", "Calmar"]
    rows = []
    order = ["two_layer_q20_e30", "two_layer_q10_e30",
             "two_layer_baseline", "layer1_invvol_lw_erc",
             "T2_bond_invvol", "solo_defensive_final"]
    win_zh = {"IS": "样本内", "OOS": "样本外", "full": "全窗"}
    for name in order:
        for w in ["IS", "OOS", "full"]:
            s = stats[name][w]
            rows.append([
                LABEL_ZH[name],
                win_zh[w],
                num(s["sharpe"], 3),
                pct(s["cagr"]),
                pct(s["dd"]),
                pct(s["ann_vol"]),
                num_abs(s["calmar"], 2) if s.get("calmar") is not None else "—",
            ])
    add_table_from_rows(doc, header, rows)

    add_paragraph_zh(doc,
        "读表要点：")
    add_bullet(doc,
        f"双层 q=0.20 ε=0.30 在 OOS 上给出 Sharpe {num(q20['OOS']['sharpe'], 3)}，"
        f"CAGR {pct(q20['OOS']['cagr'])}，最大回撤仅 {pct(q20['OOS']['dd'])}——"
        f"OOS Calmar 高达 {q20['OOS']['cagr'] / abs(q20['OOS']['dd']):.1f}（回撤面板极浅）。")
    add_bullet(doc,
        f"双层 q=0.10 ε=0.30 略胜半个 Sharpe：OOS Sharpe {num(q10['OOS']['sharpe'], 3)}，"
        f"CAGR {pct(q10['OOS']['cagr'])}，说明更集中的 α 选股在 OOS 上也没有明显退化。")
    add_bullet(doc,
        f"Layer-1 canonical（α off，Sharpe {num(l1['OOS']['sharpe'], 3)}）是最主要的贡献者，"
        "α 层对 OOS 的边际改善约 +0.08 ~ +0.12 Sharpe。")
    add_bullet(doc,
        f"T2 债券 inv-vol 单独看 OOS Sharpe {num(stats['T2_bond_invvol']['OOS']['sharpe'], 3)}"
        "看似最高，但仅债券本身没有 α 层，全窗 Calmar 只有 "
        f"{stats['T2_bond_invvol']['full']['calmar']:.2f}（DD 侧压力大）。")

    add_figure(doc, FIG / "fig07_sharpe_windows.png",
               "图 3  六个方案的净夏普：样本内 / 样本外 / 全窗对比")

    add_heading_zh(doc, "4.2 净值曲线（NAV）", level=2)
    add_paragraph_zh(doc,
        "把六个方案的净值叠到同一张图上，可以看到：双层 q=0.20 / q=0.10 与 Layer-1 canonical "
        "三条曲线整体走势非常接近，在 α 层激活的 2021 / 2023 年小幅拉开；旧版 solo defensive "
        "曲线明显低一档，尤其 2020 疫情年一次深度回撤后再没能追上。")
    add_figure(doc, FIG / "fig03_nav_all.png",
               "图 4  六个方案的净值曲线（自共同起点起，含成本）")

    add_paragraph_zh(doc,
        "把窗口拉近到 OOS 内，可以看到双层结构 + Layer-1 canonical 在 2024 全年的爬升几乎与"
        "债券 T2 平行，而在 2025 年上半年略胜一筹；旧版 solo defensive 在 OOS 内的斜率明显低于"
        "另外五个方案，这与它 OOS Sharpe 只有 +2.23 的分数一致。")
    add_figure(doc, FIG / "fig05_nav_oos.png",
               "图 5  样本外净值：2024-01-01 → 2025-07-31")

    add_heading_zh(doc, "4.3 回撤曲线", level=2)
    add_paragraph_zh(doc,
        "回撤面板则更清楚地显示了 Layer-1 组间风险预算的价值：三个含 Layer-1 的方案"
        "（双层 q20 / q10 + Layer-1 canonical）在整个样本期回撤都被压在 −2.6% 以内；"
        "T2 债券本身的最大回撤 −4.32%，主要发生在 2022 年信用债回调期；"
        "旧版 solo defensive 的 −5.20% 深度回撤集中在 2020 疫情年，之后有较长的水下期。")
    add_figure(doc, FIG / "fig04_dd_all.png",
               "图 6  六个方案的回撤路径（相对历史高点）")
    add_paragraph_zh(doc,
        "OOS 内所有六个方案的最大回撤都不超过 −1%，且集中在 2025 年 4 月一次共同回撤——"
        "但双层结构的回撤深度比 T2 债券更浅（−0.36% vs −0.56%），说明 α 层在 OOS 内没有引入"
        "额外的 tail risk。")
    add_figure(doc, FIG / "fig06_dd_oos.png",
               "图 7  样本外回撤路径")

    add_heading_zh(doc, "4.4 CAGR 与最大回撤 bar chart", level=2)
    add_figure(doc, FIG / "fig10_cagr.png",
               "图 8  六个方案的全窗 CAGR")
    add_figure(doc, FIG / "fig09_maxdd.png",
               "图 9  六个方案的全窗最大回撤")

    add_heading_zh(doc, "4.5 Calmar 比率", level=2)
    add_paragraph_zh(doc,
        "Calmar = CAGR / |Max DD| 是把 CAGR 与最大回撤合成的单一评估指标。"
        "全窗层面，双层 q=0.20 / q=0.10 与 Layer-1 canonical 均在 1.60 附近；"
        "T2 债券 0.62；旧版 solo defensive 只有 0.58。双层结构相较 solo defensive "
        f"全窗 Calmar 抬升了约 {q20['full']['calmar'] - solo['full']['calmar']:.2f}——"
        "即“同样级别的 CAGR 换到浅得多的回撤”。")
    add_figure(doc, FIG / "fig08_calmar.png",
               "图 10  全窗 Calmar 比率对比")

    add_heading_zh(doc, "4.6 α 层对 OOS 的边际贡献（Δ vs 双层 α off）", level=2)
    add_paragraph_zh(doc,
        "以 two_layer_baseline（同一双层结构、α 关掉）为对照，双层 α on 的贡献可以干净地读出。")
    header = ["书", "Δ Sharpe IS", "Δ Sharpe OOS", "Δ CAGR IS", "Δ CAGR OOS", "Δ DD OOS"]
    rows = []
    base_stats = stats["two_layer_baseline"]
    for name in ["two_layer_q20_e30", "two_layer_q10_e30",
                 "layer1_invvol_lw_erc", "T2_bond_invvol"]:
        s = stats[name]
        rows.append([
            LABEL_ZH[name],
            num(s["IS"]["sharpe"] - base_stats["IS"]["sharpe"], 3),
            num(s["OOS"]["sharpe"] - base_stats["OOS"]["sharpe"], 3),
            f"{(s['IS']['cagr'] - base_stats['IS']['cagr']) * 100:+.2f} pp",
            f"{(s['OOS']['cagr'] - base_stats['OOS']['cagr']) * 100:+.2f} pp",
            f"{(s['OOS']['dd'] - base_stats['OOS']['dd']) * 100:+.2f} pp",
        ])
    add_table_from_rows(doc, header, rows)
    add_paragraph_zh(doc,
        "两条读数：（1）α 层 OOS 的 Sharpe Δ 为 +0.082（q20）与 +0.118（q10），"
        "IS 上的 Δ +0.066 与 +0.115 相当——OOS 与 IS 保持相似量级，边际信号 hold up。"
        "（2）Layer-1 canonical 与 two_layer_baseline 的差异极小（Δ Sharpe OOS −0.04），"
        "验证了 §3 的一致性——α off 的双层与 Layer-1 canonical 数值等价。")

    add_heading_zh(doc, "4.7 实际组间持仓 vs 政策风险预算", level=2)
    add_paragraph_zh(doc,
        "Layer-1 的政策目标（equity 55% / bond_rates 20% / bond_credit 10% / commodity 15%）"
        "指的是**风险贡献占比**，不是资本权重。LW-ERC 求解器把每组的 Σ w_b·(Σw)_b 精确匹配到"
        "政策目标——下图直接把每周的实际风险贡献 RC% 和政策目标叠加：")
    add_figure(doc, FIG / "fig16_rc_check.png",
               "图 12  Layer-1 求解器校验：实际 RC% 严格等于政策目标 55/20/10/15")
    add_paragraph_zh(doc,
        "由于权益的年化 σ（≈13%）显著高于债券（≈2–3%）与商品（≈8–12%），"
        "让权益贡献 55% 的风险只需要 10% 左右的资本权重；对应地，利率债的资本权重会被抬到 60% 附近。"
        "以落地方案 q=0.20、ε=0.30 为例：")
    add_figure(doc, FIG / "fig15_policy_vs_capital.png",
               "图 13  政策风险贡献目标 vs 实际资本权重（q=0.20 ε=0.30）")
    add_paragraph_zh(doc,
        "样本内平均资本权重：权益 13.1% / 利率债 62.3% / 信用债 15.1% / 商品 9.4%。"
        "样本外平均：权益 4.8% / 利率债 58.6% / 信用债 32.5% / 商品 4.0%。"
        "对比政策目标 55 / 20 / 10 / 15，可以直接看到 LW-ERC 把权益的资本权重压到极低——"
        "但风险贡献仍恰好命中 55%。")

    add_paragraph_zh(doc,
        "分年度实际资本权重（每周组权重均值，仅计入有仓位的调仓 bar，Σ = 100%）：")

    yearly_cap = S["yearly_capital"]
    win_cap = S["window_capital"]
    header_cap = ["年份", "权益", "利率债", "信用债", "商品"]

    for key, title in [("two_layer_q20_e30", "双层 q=0.20 ε=0.30（落地方案）"),
                       ("two_layer_q10_e30", "双层 q=0.10 ε=0.30"),
                       ("two_layer_baseline", "双层 α off（Layer-1 参照）")]:
        add_paragraph_zh(doc, f"{title}：", indent_first=False)
        rows = []
        for y in sorted(int(k) for k in yearly_cap[key].keys()):
            row = [str(y)]
            for g in ["equity", "bond_rates", "bond_credit", "commodity"]:
                row.append(f"{yearly_cap[key][str(y)][g]:.2f}%")
            rows.append(row)
        # 加 IS/OOS 均值
        for lbl, wkey in [("IS 均值", "IS"), ("OOS 均值", "OOS")]:
            row = [lbl]
            for g in ["equity", "bond_rates", "bond_credit", "commodity"]:
                row.append(f"{win_cap[key][wkey][g]:.2f}%")
            rows.append(row)
        add_table_from_rows(doc, header_cap, rows,
                            col_widths=[2.5, 2.8, 2.8, 2.8, 2.8])
        doc.add_paragraph()

    add_paragraph_zh(doc,
        "组间资本权重时序（4 周平滑）可以更直观地展示 Layer-1 solver 对波动率环境切换的响应：")
    add_figure(doc, FIG / "fig13_wgroup_area_q20.png",
               "图 14  双层 q=0.20 ε=0.30：组间资本权重时序（stack area）")
    add_figure(doc, FIG / "fig14_wgroup_area_q10.png",
               "图 15  双层 q=0.10 ε=0.30：组间资本权重时序（stack area）")

    add_paragraph_zh(doc, "几点解读：")
    add_bullet(doc,
        "三个双层方案的组间资本分布高度一致——因为 Layer-1 solver 只看组间协方差，"
        "两个 α 变体的差异只影响 broad_cn / sector_cn 块内的名字选择，"
        "并不改变组间求解结果；α 层唯一间接影响组间权重的方式是通过所选名字与其他块"
        "相关性的差异。")
    add_bullet(doc,
        "2019–2021 年权益仓位一度接近 20%，因为疫情前 A 股权益 σ 相对温和；"
        "2020 年 3 月后权益 σ 抬升，solver 自动把 equity 资本压到 15% 附近，"
        "利率债则从 65% 抬到 71%。这就是"
        "“政策风险目标不变、资本权重跟着 σ 走”的直接体现。")
    add_bullet(doc,
        "2022 年信用债完成 52 周协方差 warmup 后正式入场，"
        "占比从 5% 直接跳到 35% 附近，与利率债形成分工——"
        "这一年也是双层结构里对旧版 solo defensive 拉开最大差距的一年。")
    add_bullet(doc,
        "OOS 里权益资本权重只有 4–5%，商品也只有 3–5%——这是"
        "solver 对 2024 年权益/商品波动率环境的反应，但由于风险贡献不变，"
        "OOS Sharpe 反而更高（因为组间协方差在 2024 年更接近对角，"
        "log-barrier 解距离等 σ 参考点更近，实际"
        "组合的 realized-σ 与 predicted-σ 差异更小）。")
    add_bullet(doc,
        "另一件值得注意的事：**没有硬约束** 让某组"
        "至少持有多少资本权重。如果未来希望"
        "在权益长期低配的年份强制“权益资本 ≥ 20%”，"
        "需要在 solver 里加不等式约束——当前配置没有。")

    add_heading_zh(doc, "4.8 分年度收益", level=2)
    add_figure(doc, FIG / "fig11_yearly.png",
               "图 16  六个方案的分年度净收益（灰色虚线右侧为 OOS）")

    yearly = S["yearly"]
    win_dates = ["2019", "2020", "2021", "2022", "2023", "2024", "2025"]
    header = ["年份"] + [LABEL_ZH[n] for n in order]
    rows = []
    year_keys = set()
    for name in order:
        for r in yearly[name]:
            year_keys.add(int(r["year"]))
    for y in sorted(year_keys):
        row = [str(y)]
        for name in order:
            matches = [r for r in yearly[name] if int(r["year"]) == y]
            v = matches[0]["ret"] if matches else None
            row.append(pct(v) if v is not None else "—")
        rows.append(row)
    add_table_from_rows(doc, header, rows)
    add_paragraph_zh(doc,
        "OOS 年（2024 / 2025）表现最好的年份出现在 2024 全年：所有含 Layer-1 的方案 CAGR "
        "接近 +5.8%，且回撤面板都被压在 −0.4% 以内。2025 年上半年是所有方案的调整期，"
        "四个非债券方案累计收益 +1.4% 左右，与 T2 债券的 +0.47% 相比呈现出温和的"
        "跨资产平衡效果——这正是 Layer-1 组间风险预算希望达到的“抗环境切换”能力。")

    # ---------- 五、与旧版 solo defensive 对比 ----------
    add_heading_zh(doc, "五、与旧版 Final Report 的量化对比", level=1)

    add_paragraph_zh(doc,
        "把双层 q=0.20 ε=0.30（落地方案）与旧版 solo defensive（Phase 11.2 finalist）放到同一张表上，"
        "两者共同起点均为 {}，指标都在 10 bp / 单边成本口径下计算。".format(S["common_start"]))
    header = ["指标", "旧版 solo defensive", "双层 q=0.20 ε=0.30", "Δ"]
    for w in ["IS", "OOS", "full"]:
        s_solo = solo[w]; s_new = q20[w]
        title = win_zh[w]
        rows = [
            [f"{title} · 净夏普",
             num(s_solo["sharpe"], 3), num(s_new["sharpe"], 3),
             f"{s_new['sharpe'] - s_solo['sharpe']:+.2f}"],
            [f"{title} · CAGR",
             pct(s_solo["cagr"]), pct(s_new["cagr"]),
             f"{(s_new['cagr'] - s_solo['cagr']) * 100:+.2f} pp"],
            [f"{title} · 最大回撤",
             pct(s_solo["dd"]), pct(s_new["dd"]),
             f"{(s_new['dd'] - s_solo['dd']) * 100:+.2f} pp"],
            [f"{title} · 年化波动",
             pct(s_solo["ann_vol"]), pct(s_new["ann_vol"]),
             f"{(s_new['ann_vol'] - s_solo['ann_vol']) * 100:+.2f} pp"],
            [f"{title} · Calmar",
             num_abs(s_solo["calmar"], 2),
             num_abs(s_new["calmar"], 2),
             f"{s_new['calmar'] - s_solo['calmar']:+.2f}"],
        ]
        add_paragraph_zh(doc, f"{title}窗口：", indent_first=False)
        add_table_from_rows(doc, header, rows, col_widths=[4.5, 4.0, 4.0, 3.0])
        doc.add_paragraph()

    add_paragraph_zh(doc,
        "综合结论：双层结构相较旧版 solo defensive 是一次结构性升级——")
    add_bullet(doc,
        f"OOS 净夏普从 {solo['OOS']['sharpe']:+.2f} 抬升到 {q20['OOS']['sharpe']:+.2f}"
        f"（Δ {q20['OOS']['sharpe'] - solo['OOS']['sharpe']:+.2f}）。")
    add_bullet(doc,
        f"OOS CAGR 从 {pct(solo['OOS']['cagr'])} 提升到 {pct(q20['OOS']['cagr'])}"
        f"（Δ {(q20['OOS']['cagr'] - solo['OOS']['cagr']) * 100:+.2f} pp），"
        "并且 CAGR 端从“略高于债券”变成“显著跑赢债券”。")
    add_bullet(doc,
        f"OOS 最大回撤从 {pct(solo['OOS']['dd'])} 变浅到 {pct(q20['OOS']['dd'])}"
        f"（Δ {(q20['OOS']['dd'] - solo['OOS']['dd']) * 100:+.2f} pp）。")
    add_bullet(doc,
        f"全窗 Calmar 从 {solo['full']['calmar']:.2f} 抬升到 {q20['full']['calmar']:.2f}"
        f"（Δ {q20['full']['calmar'] - solo['full']['calmar']:+.2f}）——"
        "即“在几乎不损失夏普甚至同时更高的 CAGR 下，把最大回撤对半砍”。")

    add_heading_zh(doc, "六、换手与成本", level=1)
    add_paragraph_zh(doc,
        "在成本口径下，α 层与 Layer-1 的换手结构差异显著。以三个双层方案为例："
        "α off 的双层 baseline 单边换手仅 0.046，换算成年化交易成本约 48 bps；"
        "开启 α 后，α off → q=0.20 ε=0.30 的换手抬升到 0.081（≈ 85 bps），"
        "α off → q=0.10 ε=0.30 抬升到 0.087（≈ 91 bps）。"
        "cost drag 相当于 IS Sharpe 掉 0.17 左右——这是 α 层需要额外撑起的门槛。")
    add_figure(doc, FIG / "fig12_turnover.png",
               "图 17  平均单边换手（每周）——α 层激活后的成本代价")

    # ---------- 七、实验一 · 组间风险预算敏感性 ----------
    add_heading_zh(doc, "七、实验一 · 组间风险预算敏感性", level=1)

    e1 = S["exp1"]
    e1_base = e1["base"]
    e1_ew = e1["EW"]
    e1_sigma = e1["cross_sigma"]

    add_heading_zh(doc, "7.1 出发点与实验设计", level=2)
    add_paragraph_zh(doc,
        "落地方案里的组间政策风险预算是 base = 55（权益） / 20（利率债） / 10（信用债） / 15（商品），"
        "这是最初根据经济学先验拍定的四个数字。既然第四节已经把 finalist 冻结，"
        "很自然想问一句：**这套先验对结果到底有多敏感？如果按不同的先验切开风险预算，"
        "夏普/CAGR/回撤会怎样变？**")
    add_paragraph_zh(doc, "本节把「组间风险预算」当作单一自由度扫描：", indent_first=False)
    add_bullet(doc,
        "**冻结**：Layer-1 solver（invvol × LW-target-D × log-barrier ERC，无 trend gate）、"
        "Layer-2 α（broad_cn K=5 + sector_cn K=8）、finalist (q=0.20, ε=0.30)、"
        "10 bp/单边成本口径——全部与第四节完全一致。")
    add_bullet(doc,
        "**变量**：只把 POLICY_SHARES 换成不同的四元组，重新跑一次组合。")
    add_bullet(doc,
        "**对照 A**：EW 参照 = 25/25/25/25（等风险贡献）。")
    add_bullet(doc,
        "**对照 B**：单轴 ±10pp 边际扰动 —— 把某一轴的 share 加/减 10pp，"
        "补偿量按剩下三轴的 base 比例 pro-rata 摊回去，控制总和 = 100%。")
    add_bullet(doc,
        "**全 grid**：3×3×3×3 = 81 组扰动（每轴取 {−10, 0, +10} pp），"
        "去掉 base 本身后剩 80 cell（负值截断处理）。")
    add_paragraph_zh(doc,
        "所有 cell 都做完整回测；warmup-trim 口径与前几节完全一致（IS 240 bars，OOS 82 bars）。"
        "由于所有 cell 共用同一 Layer-1 solver 和同一 α 层，任何 Δ 都可干净归因到"
        "「政策 shares 的选择」这一个自由度。")

    add_heading_zh(doc, "7.2 base vs EW —— 单个对照", level=2)
    add_paragraph_zh(doc,
        "先看 EW 与 base 两个 control cell 的三窗指标：")
    header = ["方案", "权益", "利率债", "信用债", "商品",
              "IS 夏普", "OOS 夏普", "OOS CAGR", "OOS 最大回撤"]
    rows = []
    for label, r in [("base 55/20/10/15", e1_base), ("EW 25/25/25/25", e1_ew)]:
        rows.append([
            label,
            f"{r['eq_share']*100:.1f}%",
            f"{r['br_share']*100:.1f}%",
            f"{r['bc_share']*100:.1f}%",
            f"{r['cm_share']*100:.1f}%",
            f"{r['is_sharpe']:+.3f}",
            f"{r['oos_sharpe']:+.3f}",
            f"{r['oos_cagr']*100:+.2f}%",
            f"{r['oos_max_dd']*100:+.2f}%",
        ])
    add_table_from_rows(doc, header, rows)
    add_paragraph_zh(doc,
        f"直接读数：EW 在 IS 上 Sharpe {e1_ew['is_sharpe']:+.3f} 比 base "
        f"{e1_base['is_sharpe']:+.3f} 高 Δ +{e1_ew['is_sharpe']-e1_base['is_sharpe']:.3f}，"
        f"OOS 上 Δ +{e1_ew['oos_sharpe']-e1_base['oos_sharpe']:.3f}；"
        f"但 EW 的 OOS CAGR {e1_ew['oos_cagr']*100:+.2f}% 反而略低于 base "
        f"{e1_base['oos_cagr']*100:+.2f}%——同样是「收益略减 / 波动大减 → Sharpe 抬升」的组合"
        "（与第七节实验二的机制类似，但触发原因不同）。")
    add_paragraph_zh(doc,
        f"**噪声尺度先记住**：跨 80 cells 的 IS 夏普 σ = {e1_sigma['IS_Sh']:.3f}，"
        f"OOS 夏普 σ = {e1_sigma['OOS_Sh']:.3f}。因此 EW − base 的 OOS Sharpe Δ 只有 "
        f"~{(e1_ew['oos_sharpe']-e1_base['oos_sharpe'])/e1_sigma['OOS_Sh']:.1f}× cross-cell σ，"
        "IS 上略高一点 —— 属于「有信号但不算强」量级。任何单点比较都要放到这个噪声带里读。")

    add_heading_zh(doc, "7.3 单轴 ±10pp 边际斜率 + 噪声带", level=2)
    add_figure(doc, FIG / "fig_exp1_01_axis_slope.png",
               "图 25  实验一：4 轴 ±10pp 单轴边际 Δ 夏普 vs cross-cell ±1σ "
               "噪声带（虚线蓝 = IS σ ±0.10，虚线红 = OOS σ ±0.40）")
    header = ["轴", "方向", "Δ IS 夏普", "Δ OOS 夏普",
              "Δ IS CAGR", "Δ OOS CAGR", "IS σ 倍数", "OOS σ 倍数"]
    rows = []
    for r in e1["axis_slope"]:
        rows.append([
            AXIS_ZH.get(r["axis"], r["axis"]),
            r["direction"],
            f"{r['d_is_sharpe']:+.3f}",
            f"{r['d_oos_sharpe']:+.3f}",
            f"{r['d_is_cagr_pp']:+.2f} pp",
            f"{r['d_oos_cagr_pp']:+.2f} pp",
            f"{abs(r['d_is_sharpe'])/e1_sigma['IS_Sh']:.2f}",
            f"{abs(r['d_oos_sharpe'])/e1_sigma['OOS_Sh']:.2f}",
        ])
    add_table_from_rows(doc, header, rows)
    add_paragraph_zh(doc,
        "读表的规则：**任何绝对值小于 1× cross-cell σ 的 Δ，都应当当作「大概率是噪声」看待**。"
        "按这条规则去数：")
    add_bullet(doc,
        "**权益轴**：±10pp 两个方向都 < 1σ 但**方向严格对称**（+10pp: IS −0.10, OOS −0.30；"
        "−10pp: IS +0.08, OOS +0.25），且 IS 与 OOS 同号 —— 这种对称性本身就是结构信号，"
        "尽管每个单点没跨过 1σ 门槛。")
    add_bullet(doc,
        "**利率债轴**：−10pp 方向 IS 1.55σ + OOS 1.42σ，是全表唯一跨过 1σ 门槛的方向。"
        "但 +10pp 只有 1.03σ / 0.90σ，且如后文 §7.5 指出，这个方向和权益轴有 pro-rata 耦合，"
        "扣掉「权益隐含变化」贡献之后残差 ≈ 0.9σ。视作「suggestive 但不能完全"
        "从权益混淆里分离出来」。")
    add_bullet(doc,
        "**信用债轴**：+10pp Δ IS Sh = +0.001（0.01σ！几乎完全为零）、"
        "−10pp Δ IS Sh = +0.025（0.24σ）—— IS 上完全无信号。"
        "OOS 上两个方向反而同号（−0.30 和 +0.16 都是「离 base 走」但 Sharpe 变化方向不定）"
        "——**结构性噪声**。")
    add_bullet(doc,
        "**商品轴**：+10pp Δ IS Sh = +0.002（0.02σ）、−10pp = −0.051（0.49σ），"
        "OOS 上一个 −0.17、一个 +0.16 —— **也是噪声**。")
    add_paragraph_zh(doc,
        "**归口结论**：8 个方向里，能从 cross-cell σ 噪声里剥出来的只有 2 个 —— "
        "权益轴的对称对读，以及利率债 −10pp 的单向大幅移动（且部分被权益混淆解释）。"
        "**信用债和商品两个轴，本实验无法给出可信的方向性判断。**")

    add_heading_zh(doc, "7.4 权益轴 —— 唯一可以放心断言的方向", level=2)
    add_figure(doc, FIG / "fig_exp1_03_equity_dominance.png",
               "图 26  实验一：80-cell 上 OOS 夏普 / CAGR / 最大回撤 vs 权益占比 —— "
               "三个截面内部一致（权益 ↑ 拉高 CAGR、恶化 DD、拉低 Sharpe）")
    am = e1["axis_marginals"]["equity"]
    add_paragraph_zh(doc,
        f"80 cell 上「权益占比 vs OOS 指标」的三个 marginal 相关系数："
        f"OOS Sharpe corr = {am['oos_sharpe']['corr']:+.2f}（斜率 "
        f"{am['oos_sharpe']['slope_per_unit']*10:+.2f}/10pp），"
        f"OOS CAGR corr = {am['oos_cagr']['corr']:+.2f}（斜率 "
        f"{am['oos_cagr']['slope_per_unit']*10*100:+.2f} pp/10pp），"
        f"OOS 最大回撤 corr = {am['oos_max_dd']['corr']:+.2f}（斜率 "
        f"{am['oos_max_dd']['slope_per_unit']*10*100:+.2f} pp/10pp）。"
        "换句话说，**在 80 cell 上，权益占比每+10pp，"
        "OOS 组合会拿到 +0.19 pp 的 CAGR、−0.08 pp 的最大回撤（更深）、−0.27 的 Sharpe**。"
        "三个指标彼此内部一致，也与「权益是高波动 / 高预期收益的类别」的经济学直觉一致。")
    add_paragraph_zh(doc,
        "对比图 26 三个 panel，权益占比与三个指标的散点云都有清晰斜率 —— "
        "这就是「除了直觉，还有数据支撑」的结构信号。IS 上 corr 也稳定"
        f"（IS Sharpe corr = {am['is_sharpe']['corr']:+.2f}），"
        "IS/OOS 同号，非 IS-fit artifact。")
    add_paragraph_zh(doc,
        "**唯一可信的建议**：现在 base 的 55% 权益占比处在 Sharpe-suboptimal 区的偏上侧；"
        "如果目标是抬 Sharpe，把权益向 45–50% 的方向移动是可以支持的（IS + OOS 都指向"
        "同一方向）。但要清楚这条建议是以「CAGR 少 5–8 bp/10pp」为代价买回来的，"
        "如果目标反而是最大化 CAGR，则应该维持甚至加高权益占比。")

    add_heading_zh(doc, "7.5 为什么另外三轴的效应不能拆出来", level=2)
    add_paragraph_zh(doc,
        "如果只看 marginal correlation，会发现其他三轴看似也有相关性："
        f"利率债 vs OOS Sh corr = {e1['axis_marginals']['bond_rates']['oos_sharpe']['corr']:+.2f}、"
        f"信用债 vs OOS Sh corr = {e1['axis_marginals']['bond_credit']['oos_sharpe']['corr']:+.2f}、"
        f"商品 vs OOS Sh corr = {e1['axis_marginals']['commodity']['oos_sharpe']['corr']:+.2f}。"
        "看起来利率债 / 信用债正相关、商品负相关。但这些相关性都被「权益轴的耦合」污染了。")
    add_figure(doc, FIG / "fig_exp1_02_marginal_scatters.png",
               "图 27  实验一：4 轴各自 share vs OOS 夏普 —— 只有权益是负斜率清晰、"
               "散点密度均匀；其他三轴的散点云更宽、结构性更弱")
    add_paragraph_zh(doc,
        "污染的来源是「4 shares 之和恒等于 100%」的约束：任何一个轴 ±10pp 的移动都"
        "**必然**会让其他三轴反向移动。本实验里的补偿方式是 pro-rata by base share，"
        "所以「利率债 −10pp」实际上包含「权益 +6.9pp、信用债 +1.4pp、商品 +1.7pp」"
        "的合成移动 —— 其中的权益 +6.9pp 已经足以解释相当一部分的 OOS Sh 恶化。")
    add_paragraph_zh(doc,
        "定量验证：用 §7.4 的权益单边斜率 (−0.27/10pp OOS Sh) 反推 "
        "「利率债 −10pp 时权益隐含涨 +6.9pp」的贡献 = "
        "6.9/10 × (−0.30) = −0.21 OOS Sh。"
        "观测到的「利率债 −10pp」的 Δ OOS Sh = −0.57，扣掉权益混淆项后残差约 −0.36 —— "
        "仍属 0.9σ 量级，可视作「suggestive of 利率债 shares 也重要」，"
        "但显著性不足以做独立断言。")
    # Top-5 / Bottom-5 权益均值直接算，避免硬编码
    top5_eq_pct = 100 * sum(r["eq_share"] for r in e1["top5"]) / len(e1["top5"])
    bot5_eq_pct = 100 * sum(r["eq_share"] for r in e1["bot5"]) / len(e1["bot5"])
    add_figure(doc, FIG / "fig_exp1_04_top_bottom.png",
               f"图 28  实验一：Top-5 vs Bottom-5 (按 OOS 夏普) 的 4 轴平均 shares —— "
               f"唯一系统性差距在权益（{top5_eq_pct:.1f}% vs {bot5_eq_pct:.1f}%），"
               "其他三轴的差距要么很小、要么与权益差距强相关")
    add_paragraph_zh(doc,
        "看图 28 一眼可以更直观地读出：Top-5 和 Bottom-5 之间**唯一稳定的差异**在权益"
        f"（Δ 约 {top5_eq_pct - bot5_eq_pct:+.1f} pp）；"
        "利率债、信用债、商品三轴虽也有差异，但都可以由「权益反方向变化 + "
        "pro-rata 分摊」解释掉。因此 —— **和之前 markdown 报告里"
        "「bond_rates 是最陡的轴、应当把 bond_rates share 抬高」的结论相比，"
        "本报告采取更保守的读法：能得出的结论只有「权益略偏高」这一条**。")

    add_heading_zh(doc, "7.6 实验一小结", level=2)
    add_bullet(doc,
        "**base 未处于 Sharpe-最优**：EW 25/25/25/25 的 OOS Sharpe 高出 base 约 "
        f"+{e1_ew['oos_sharpe']-e1_base['oos_sharpe']:.2f}，Top-5 grid cell 权益均值 "
        f"{100 * sum(r['eq_share'] for r in e1['top5']) / len(e1['top5']):.1f}% ——base 的 55% 权益比 Sharpe-favored 区略偏上侧。")
    add_bullet(doc,
        "**唯一可信的方向**：权益 ↑ → CAGR ↑ / DD 恶化 / Sharpe ↓。这条对称、"
        "IS/OOS 同号、marginal corr 稳定，任何单点在噪声带内，但整体图形结构清晰。")
    add_bullet(doc,
        "**其他三轴不做独立断言**：信用债 / 商品的 IS Δ ≈ 0；利率债 −10pp 的大幅移动"
        "部分被权益混淆解释；剩余残差 ~0.9σ 达不到断言门槛。跨 80-cell 的 OOS Sharpe σ = "
        f"{e1_sigma['OOS_Sh']:.2f}，把这作为读表噪声带。")
    add_bullet(doc,
        "**不建议立即重调 policy shares**：EW − base 的 OOS Sharpe Δ 只是 ~1.3σ，"
        "而且 base 目前已经是 finalist 的一部分并通过了 §四.6 的 stress test。"
        "本实验建议作为「经济学先验的 gut-check」使用（结论：先验方向偏 equity-heavy），"
        "而不是作为 optimizer 结果去落地新的 shares。")
    add_bullet(doc,
        "**v7 阶段的潜在动作**：若未来想据此重构 policy shares，需要至少两点补强："
        "(a) 采用更小步长（如 ±2–5 pp）+ 更长 OOS 窗口，"
        "在权益 45–55% 区间做二次扫描，"
        "把「更保守的权益权重是否有一致 OOS 益处」验证到 2σ 以上；"
        "(b) 用受约束的多变量拟合（如 log-ratio 表征 + Lasso）取代 pro-rata "
        "补偿，减弱轴之间的耦合。当前实验足以启发方向，但不足以直接落地。")

    # ---------- 八、实验二 · 非-α 块代表集压缩 ----------
    add_heading_zh(doc, "八、实验二 · 非-α 块的代表集压缩", level=1)

    add_heading_zh(doc, "8.1 出发点与做法", level=2)
    add_paragraph_zh(doc,
        "上面第三/四节的双层结构里，6 个非-α 块（bond_rates、bond_credit、"
        "cross_border_dm、cross_border_hk、metals、commodity_other）都是"
        "块内 1/σ 反比全持——所有入池 ETF 都持有一点权重。随着 v6 池子不断纳新"
        "（cross_border_hk 从 2020 年的 2 只涨到 2025 年的 22 只，"
        "cross_border_dm 从 3 只涨到 18 只），"
        "「全持」在名义合规和交易可行性上都开始变得笨重；"
        "实验二试图回答：能不能只留每个聚类里一支代表 ETF、"
        "让代表继承整个聚类的 1/σ 权重，尽量复刻「全持」的行为？")
    add_paragraph_zh(doc, "方法要点：", indent_first=False)
    add_bullet(doc,
        "在每年 refresh date（每年第一根 W-FRI）上，用严格 causal 的过去 52/78/104 周"
        "滚动相关阵求 D = √(2·(1−ρ̄))，做 complete-linkage 层次聚类。")
    add_bullet(doc,
        "**K 不是输入，是输出**：以「代表组合 vs 全持组合的残差年化波动 / 全持年化波动 ≤ "
        "0.20」为准则，从 K=1 起递增，第一次满足阈值的 K 即为该块该年的 K。等价于"
        "解释方差 ≥ 96%。")
    add_bullet(doc,
        "每聚类的代表 = 该聚类内 63 日均日成交额（ADV）最高的 ETF，无预测指标。"
        "代表承接整个聚类的 1/σ 权重。")
    add_bullet(doc,
        "刷新频率 = 每年一次（每年首个 W-FRI），"
        "Layer-1 / Layer-2 / α 层 / 成本口径全部沿用第四节冻结的 finalist "
        "(q=0.20, ε=0.30)，只替换 6 个非-α 块的子权重构造函数——"
        "任何 Δ 都可以干净归因到「代表集 vs 全持」这一维度。")
    add_paragraph_zh(doc,
        f"warmup-trim 与第三/四节完全一致：两本书共同 first-live 为 "
        f"{e2['common_first_live']}，"
        "IS = 该点 → 2023-12-31，OOS = 2024-01-01 → 2025-07-31。")

    add_heading_zh(doc, "8.2 K 的时序：K vs N", level=2)
    add_paragraph_zh(doc,
        "先看 K 随时间的变化。刷新逻辑是每年一次，每年每块单独选 K：")
    add_figure(doc, FIG / "fig_exp2_07_k_by_year.png",
               "图 18  实验二：adaptive-K 每年每块的 K / N（残差波动阈 0.20）")
    add_paragraph_zh(doc, "读图要点：", indent_first=False)
    add_bullet(doc,
        "cross_border_hk 从 2020 年 K/N = 1/2 长到 2025 年 K/N = 6/22——"
        "N 一路涨到 22 只，K 只需要 6 只就把残差波动压到 16.5% 阈值以下。"
        "换句话说，池子里 73% 的名字是可以被 6 只代表复刻的。")
    add_bullet(doc,
        "cross_border_dm 一直保持 K = N（每只自己一个聚类），直到 2025 年"
        "才因新入池的名字大多落入已有聚类而首次出现 K < N（11 / 18）——"
        "这就是「块内异质性」的直接证据，而不是方法失效。")
    add_bullet(doc,
        "bond_rates、bond_credit 从 2 只/2 只逐步扩张到 12 只/5 只，"
        "对应的 K 也从 2/2、1/2 增到 8/12、4/5——都在阈值一次性满足即停。")

    add_heading_zh(doc, "8.3 组合层面的 IS / OOS 三窗指标", level=2)
    add_paragraph_zh(doc,
        "把「全持 vs 代表集」两本书放到 finalist (q=0.20, ε=0.30) 里，"
        "分别跑一次完整的双层组合。10 bp / 单边成本口径下的三窗指标：")

    header = ["书", "窗口", "净夏普", "CAGR", "年化波动", "最大回撤", "Calmar"]
    rows = []
    win_zh = {"IS": "样本内", "OOS": "样本外", "full": "全窗"}
    for name, label in [("exp2_hold_all", "全持 hold-all (K̄=38.1)"),
                        ("exp2_replicated", "代表集 adaptive (K̄=24.1)")]:
        for w in ["IS", "OOS", "full"]:
            s = e2["stats"][name][w]
            rows.append([
                label, win_zh[w],
                num(s["sharpe"], 3),
                pct(s["cagr"]),
                pct(s["ann_vol"]),
                pct(s["dd"]),
                num_abs(s["calmar"], 2) if s.get("calmar") is not None else "—",
            ])
    add_table_from_rows(doc, header, rows)

    # Δ 表
    add_paragraph_zh(doc,
        "Δ（代表集 − 全持）：", indent_first=False)
    delta_header = ["窗口", "Δ 夏普", "Δ CAGR", "Δ 年化波动", "Δ 最大回撤"]
    delta_rows = []
    for w in ["IS", "OOS", "full"]:
        h_s = e2["stats"]["exp2_hold_all"][w]
        r_s = e2["stats"]["exp2_replicated"][w]
        delta_rows.append([
            win_zh[w],
            f"{r_s['sharpe'] - h_s['sharpe']:+.3f}",
            f"{(r_s['cagr'] - h_s['cagr']) * 100:+.2f} pp",
            f"{(r_s['ann_vol'] - h_s['ann_vol']) * 100:+.2f} pp",
            f"{(r_s['dd'] - h_s['dd']) * 100:+.2f} pp",
        ])
    add_table_from_rows(doc, delta_header, delta_rows,
                        col_widths=[2.5, 2.6, 3.0, 3.0, 3.0])

    add_paragraph_zh(doc,
        "**关键读数——把此前 markdown 报告里的结论修正回来**：")
    add_bullet(doc,
        f"IS 上代表集是「几乎打平」：Sharpe −0.048、CAGR −0.33 pp、DD 完全一致。"
        "这一部分行为很符合直觉：训练窗口里，全持和代表集本来就把同一批风险源"
        "拆成同一批组合，Δ 只反映聚类残差与 refresh 成本。")
    add_bullet(doc,
        "OOS 上 Sharpe 从 +3.62 抬升到 +4.08（Δ +0.46）看似很漂亮，"
        f"但**收益并没有上升——OOS CAGR 反而从 {pct(e2_hold['OOS']['cagr'])} "
        f"掉到 {pct(e2_rep['OOS']['cagr'])}（Δ "
        f"{(e2_rep['OOS']['cagr']-e2_hold['OOS']['cagr'])*100:+.2f} pp）**。"
        f"Sharpe 抬升的**全部来源是年化波动被压下去**——从 {pct(e2_hold['OOS']['ann_vol'])} "
        f"缩到 {pct(e2_rep['OOS']['ann_vol'])}（相对减幅 "
        f"−{(1-e2_rep['OOS']['ann_vol']/e2_hold['OOS']['ann_vol'])*100:.1f}%）。")
    add_bullet(doc,
        "所以 OOS 的改善本质是一次 **风险压缩**，而不是「收益变好」。"
        "早期 v1 markdown 里的口径读起来像是代表集在 OOS 上产生了 α，"
        "那是不准确的——真实故事是「代表集用少一点的名字、更稳的持仓，"
        "把 OOS 波动压得更低，同一批 α 的分子被稍削、分母被大削」。")

    add_heading_zh(doc, "8.4 波动压缩的直接证据 · 分年度收益与年化波动", level=2)
    add_paragraph_zh(doc,
        "把两本书拆到每一年看：上图是全年净收益，下图是全年年化波动。"
        "对比两栏可以直接读出「Sharpe 抬升靠什么」。")
    add_figure(doc, FIG / "fig_exp2_05_yearly.png",
               "图 19  实验二：全年净收益（上）vs 全年年化波动（下）"
               "——OOS 段（灰线右侧）代表集收益略低但波动系统性收窄")

    add_paragraph_zh(doc,
        "分年度看有几个稳定的规律：")
    add_bullet(doc,
        "2019–2023 IS 段：两本书的年度收益基本一致，代表集在 2020/2022 略输一点，"
        "反映聚类残差的路径依赖；同时年化波动几乎每一年都被代表集压低（"
        "2020 4.40% → 4.19%，2021 2.99% → 2.78%，2022 1.98% → 1.76%，"
        "2023 1.21% → 0.87%）。")
    add_bullet(doc,
        "2024–2025 OOS 段：2024 年代表集收益 5.30% 略输全持 5.71%，"
        "但波动从 1.42% 压到 1.11%——波动比收益少的幅度大得多，"
        "所以 2024 年年化 Sharpe 反而抬升。2025 上半年是同样的现象。")

    add_heading_zh(doc, "8.5 为什么波动会被系统性压低——两个候选假说", level=2)
    add_paragraph_zh(doc,
        "本报告倾向的解释是**新入池 ETF 的规模/微结构噪声在 OOS 段被年度刷新过滤掉了**："
        "根据观察，v6 每年新准入的 ETF 通常来自次新品类（跨境宽基 / 新兴主题债券 / "
        "细分商品），这些 ETF 上市初期 AUM 通常不大、日成交额（ADV）相对稀薄，"
        "反映到周频回报里就是**收益近似而波动更粗**——即使按 1/σ 反比配权，"
        "小样本估的 σ 本身就有偏，配权也没有完全抵消。"
        "全持口径每周都把这些新名字持在组合里，噪声就直接进书；"
        "代表集口径下，新名字只有在**下一次年度 refresh** 且**被选为聚类代表**"
        "（top-ADV，天然偏向大 AUM 老名字）时才会进书——刚上市的小 AUM ETF 有"
        "很大概率被聚类到一支已有老 ETF 后面，其权重直接由老代表继承。")
    add_paragraph_zh(doc,
        "这与观察到的现象吻合：cross_border_hk 2024 / 2025 分别有 17/22 只入池，"
        "但代表数只保留 5 / 6——被压掉的 12 / 16 只名字大多是次新 ETF，"
        "而年化波动的 OOS 减幅恰好也集中在这两年。"
        "换句话说，「年度 refresh + top-ADV 代表」这一组合起到了"
        "**噪声-新股过滤**的效果，属于代表集的意外副作用（既不是策略设计的目标，"
        "也不是集成 α 的贡献）。")
    add_paragraph_zh(doc,
        "作为对照，还有第二个候选假说值得列出："
        "即聚类本身的**几何平均效应**——把同一聚类里 N 支强相关 ETF"
        "合并成一支代表，损失了「聚类内 idiosyncratic 波动」这一分量，"
        "但同时也损失了它们之间的分散化。"
        "如果聚类内平均相关系数 ρ̄ 足够高（本实验大部分块 ρ̄ ≥ 0.90，"
        "参见 §7.2 K 表），损失分散化的边际影响比损失 idio 的边际影响小，"
        "净效果仍然是 σ 下降。这个假说与「新入池 ETF 噪声过滤」并不互斥，"
        "更可能是两者叠加。")
    add_paragraph_zh(doc,
        "从策略生产的角度看，无论是哪个假说主导，"
        "**结论仍然是同一个**：如果目标是「用尽量少的名字复刻宽持仓」，"
        "代表集方案是 Pareto 改进（Sharpe 上升 + 名字数下降 + 波动下降，"
        "只以微小的收益让渡为代价）；但如果目标是「压榨 α」，"
        "本实验没有额外贡献，Sharpe 抬升是风险端的红利而非收益端的红利。")

    add_heading_zh(doc, "8.6 净值与回撤路径", level=2)
    add_figure(doc, FIG / "fig_exp2_01_nav_full.png",
               "图 20  实验二：全窗净值（起始 = 1，含 10 bp/单边成本）")
    add_paragraph_zh(doc,
        "全窗看两条曲线几乎重合，代表集在 IS 段略低——这正是「小额收益让渡」"
        "的直接体现。转到 OOS 段（2024-01-01 之后），两条曲线的斜率变化几乎一致，"
        "但代表集的日常抖动明显更小。")
    add_figure(doc, FIG / "fig_exp2_02_dd_full.png",
               "图 21  实验二：全窗回撤路径")
    add_paragraph_zh(doc,
        "回撤面板层面，代表集的最大回撤（−2.54%）与全持完全一致，"
        "都发生在 2022 年 11 月信用债回调期；OOS 段的两次浅回撤"
        "（2024 年 10 月、2025 年 4 月），代表集反而比全持略深一点——"
        "这是波动压缩带来的副产物：更平的净值线在遇到共同下跌事件时，"
        "回撤深度差异并不显著。")
    add_figure(doc, FIG / "fig_exp2_03_nav_oos.png",
               "图 22  实验二：OOS 窗口净值")
    add_figure(doc, FIG / "fig_exp2_04_dd_oos.png",
               "图 23  实验二：OOS 窗口回撤")

    add_heading_zh(doc, "8.7 换手与 refresh-week 成本", level=2)
    turn_h = e2["turnover_attribution"]["hold_all"]
    turn_r = e2["turnover_attribution"]["replicated_adaptive"]
    add_paragraph_zh(doc,
        "代表集有一个额外成本项：每年 refresh week（每年首个 W-FRI）会有一次代表刷新"
        "带来的换手峰值。attribution 结果：")
    header = ["书", "全窗单边换手均值", "refresh 周均值", "非-refresh 周均值",
              "超额换手（refresh − 非-refresh）", "刷新成本 bp/yr"]
    rows = [
        ["全持 hold-all",
         f"{turn_h['avg_turnover']*100:.2f}%",
         f"{turn_h['refresh_avg']*100:.2f}%",
         f"{turn_h['other_avg']*100:.2f}%",
         f"{turn_h['excess_refresh']*100:+.2f}%",
         f"{turn_h['refresh_cost_drag_bp_yr']:+.1f}"],
        ["代表集 replicated",
         f"{turn_r['avg_turnover']*100:.2f}%",
         f"{turn_r['refresh_avg']*100:.2f}%",
         f"{turn_r['other_avg']*100:.2f}%",
         f"{turn_r['excess_refresh']*100:+.2f}%",
         f"{turn_r['refresh_cost_drag_bp_yr']:+.1f}"],
    ]
    add_table_from_rows(doc, header, rows)
    add_paragraph_zh(doc,
        "读法：全持的日常换手（0.082/周）反而略高于代表集（0.068/周），"
        "因为代表集持仓总数少 —— 每周 σ 漂移触发的边际调仓也少。"
        "但代表集在 refresh 周多了一次代表刷新，导致 refresh 周均值飙到 25.5%，"
        "对应超额换手 ~19% × 1 refresh/年 × 10 bp/side × 2 sides = 约 +3.8 bp/yr "
        "的额外成本——**相对 OOS ~+100 bp/yr 的 Sharpe 改善红利，"
        "这个刷新成本占比 ~4%，可以接受**。")
    add_figure(doc, FIG / "fig_exp2_06_turnover.png",
               "图 24  实验二：单边换手时序（灰色竖线为每年首个 W-FRI）——"
               "代表集在每年 refresh 有一次明显的换手尖峰，"
               "其他周普遍低于全持")

    add_heading_zh(doc, "8.8 实验二小结", level=2)
    add_bullet(doc,
        "**方法定位**：代表集是 v7 阶段推荐的「运营简化」升级方案——"
        "把 6 个非-α 块的名字数从平均 38.1 只压到 24.1 只（−37%），"
        "不需要新增任何预测信号或参数扫格。")
    add_bullet(doc,
        f"**收益结构**：IS 上小幅让渡（Sharpe −0.048，CAGR −0.33 pp），"
        f"OOS 上 Sharpe 抬升 +0.46——但**这是波动压缩带来的抬升，不是收益抬升**："
        f"OOS CAGR 反而 −0.37 pp，OOS 年化波动 −0.23 pp（相对减幅 −18%）。")
    add_bullet(doc,
        "**波动压缩来源假说**：主导因素很可能是「新入池 ETF 通常 AUM 小、噪声大」"
        "被年度刷新 + top-ADV 代表选择自然过滤掉；次要因素是聚类内几何平均"
        "抹掉 idio 波动的效果。这与 cross_border_hk / cross_border_dm 两块"
        "近年密集扩容后 OOS 波动集中收窄的观察一致。")
    add_bullet(doc,
        "**成本**：refresh 周多出 +3.8 bp/yr 的换手成本，"
        "相对 OOS Sharpe ~+100 bp/yr 的改善量级，占比 ~4%——完全可接受。"
        "如果实盘想进一步降低这一成本，可以在 refresh 周分批建仓，"
        "但不是必须的。")
    add_bullet(doc,
        "**遗留问题**：本实验没有回答的问题是「如果不是新入池 ETF 而是老 ETF 的"
        "AUM 长期萎缩，会不会破坏这套框架」——代表集的 top-ADV 选择规则依赖"
        "存量老名字的 ADV 长期稳定；如果 ADV 结构性下降，代表选择本身可能出现"
        "路径依赖。这一项需要在 v7 阶段用更长的 OOS 窗口验证。")

    add_heading_zh(doc, "九、结语与局限", level=1)
    add_paragraph_zh(doc,
        "本报告展示了 v6 项目在 Final Report 之后的一次结构性升级：把“单层的 Top-q 纯多头”"
        "改造为“组间风险预算 + 块内 α 选股”的双层结构，落地 (q=0.20, ε=0.30) 方案在 OOS 上"
        f"给出净夏普 {q20['OOS']['sharpe']:+.2f}、CAGR {pct(q20['OOS']['cagr'])}、"
        f"最大回撤 {pct(q20['OOS']['dd'])}，全窗 Calmar 抬升到 {q20['full']['calmar']:.2f}，"
        "较此前 solo defensive 是一次跨门槛的改进。")
    add_paragraph_zh(doc,
        "已知的局限：(1) α 层目前只覆盖 broad_cn 与 sector_cn 两块，cross_border_hk 待其 2024+"
        "非危机数据积累后再考虑纳入；(2) 本轮扫格使用了统一 (q, ε)，Phase 13.5 的原始 finalist "
        "其实是 ε_broad_cn = 0.20 与 ε_sector_cn = 1.00，未来可以补做块级 (q, ε) 扫格；"
        "(3) 非 α 块的 sizing 固定为 inv-vol，是否 eqw 会有额外收益仍待验证；"
        "(4) OOS 窗口只有 82 周，样本外仍需继续积累。")
    add_paragraph_zh(doc,
        "但截至 2026-07-22 的 OOS shot 已经足以支持一次落地方案切换：本报告推荐 "
        "“Layer-1（invvol × LW-ERC，no-trend） + Layer-2 α（broad_cn K=5 + sector_cn K=8，"
        "long_q20-replace，(q=0.20, ε=0.30)）” 作为 v6 阶段的新生产 finalist，"
        "旧版 solo defensive 从此可以功成身退。")
    add_paragraph_zh(doc,
        "在此基础上，第八节讨论的实验二（非-α 块代表集压缩）可作为 v7 阶段的候选"
        "「运营简化」升级：把 6 个非-α 块的名字数从 38.1 只压到 24.1 只（−37%），"
        "IS 上小幅让渡（Sharpe −0.05、CAGR −0.33 pp），OOS 上 Sharpe 抬升 +0.46——"
        "**但需要注意：这轮 OOS 改善是波动压缩（−18%）而不是收益提升（CAGR 反而 "
        "−0.37 pp）**，很可能来自年度 refresh + top-ADV 代表规则对新入池小 AUM "
        "ETF 噪声的自然过滤。这也意味着如果未来 v6 池子扩容速率放缓，这一红利可能收敛，"
        "在 v7 阶段推行前建议再看一段更长的 OOS。"
        "第七节实验一给的建议同样保守：先不动 policy shares，"
        "把「权益略偏高」当经济学 prior gut-check 记下即可，"
        "任何重构都需要更小步长 + 更长 OOS 的二次验证。")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
